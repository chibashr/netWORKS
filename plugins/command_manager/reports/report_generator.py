#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Report generator for the Command Manager plugin
"""

import datetime
from pathlib import Path
from loguru import logger

from PySide6.QtCore import Qt
from PySide6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, 
    QTableWidget, QTableWidgetItem, QHeaderView, QMessageBox,
    QFileDialog, QGroupBox, QFormLayout, QCheckBox,
    QDialogButtonBox, QLineEdit, QComboBox
)

class ReportGenerator(QDialog):
    """Dialog for generating command output reports"""
    
    def __init__(self, plugin, parent=None):
        """Initialize the dialog"""
        super().__init__(parent)
        
        self.plugin = plugin
        
        # Set dialog properties
        self.setWindowTitle("Generate Command Report")
        self.resize(700, 500)
        
        # Create UI components
        self._create_ui()
        
        # Load devices
        self._load_devices()
        
    def _create_ui(self):
        """Create the UI components"""
        # Main layout
        layout = QVBoxLayout(self)
        
        # Device selection section
        device_group = QGroupBox("Select Devices")
        device_layout = QVBoxLayout(device_group)
        
        self.device_table = QTableWidget()
        self.device_table.setColumnCount(2)
        self.device_table.setHorizontalHeaderLabels(["Device", "IP Address"])
        self.device_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.device_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.device_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.device_table.setSelectionMode(QTableWidget.MultiSelection)
        
        device_layout.addWidget(self.device_table)
        
        # Report options section
        options_group = QGroupBox("Report Options")
        options_layout = QFormLayout(options_group)
        
        # Report format
        self.report_format = QComboBox()
        self.report_format.addItems(["Text (.txt)", "HTML (.html)", "Excel (.xlsx)", "Word (.docx)"])
        options_layout.addRow("Format:", self.report_format)
        
        # Report title
        self.report_title = QLineEdit("Command Output Report")
        options_layout.addRow("Title:", self.report_title)
        
        # Include device info
        self.include_device_info = QCheckBox("Include device information")
        self.include_device_info.setChecked(True)
        options_layout.addRow("", self.include_device_info)
        
        # Include all commands
        self.include_all_commands = QCheckBox("Include all command outputs")
        self.include_all_commands.setChecked(True)
        options_layout.addRow("", self.include_all_commands)
        
        # Date range
        date_layout = QHBoxLayout()
        
        self.date_from = QLineEdit()
        self.date_from.setPlaceholderText("YYYY-MM-DD")
        
        self.date_to = QLineEdit()
        self.date_to.setPlaceholderText("YYYY-MM-DD")
        
        date_layout.addWidget(QLabel("From:"))
        date_layout.addWidget(self.date_from)
        date_layout.addWidget(QLabel("To:"))
        date_layout.addWidget(self.date_to)
        
        options_layout.addRow("Date Range:", date_layout)
        
        # Include success only
        self.success_only = QCheckBox("Include only successful commands")
        self.success_only.setChecked(False)
        options_layout.addRow("", self.success_only)
        
        # Add groups to main layout
        layout.addWidget(device_group)
        layout.addWidget(options_group)
        
        # Buttons
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self._on_generate)
        buttons.rejected.connect(self.reject)
        
        layout.addWidget(buttons)
        
    def _load_devices(self):
        """Load device list"""
        # Get all devices
        devices = self.plugin.device_manager.get_devices()
        
        # Add devices to table
        for device in devices:
            alias = device.get_property("alias", "Unnamed Device")
            ip_address = device.get_property("ip_address", "")
            
            # Skip devices without command outputs
            outputs = self.plugin.get_command_outputs(device.id)
            if not outputs:
                continue
                
            row = self.device_table.rowCount()
            self.device_table.insertRow(row)
            
            # Device info
            alias_item = QTableWidgetItem(alias)
            alias_item.setData(Qt.UserRole, device.id)
            
            ip_item = QTableWidgetItem(ip_address)
            
            # Add to table
            self.device_table.setItem(row, 0, alias_item)
            self.device_table.setItem(row, 1, ip_item)
            
    def _on_generate(self):
        """Handle generate button"""
        # Get selected devices
        selected_devices = []
        for item in self.device_table.selectedItems():
            # Make sure we only count each row once
            if item.column() == 0:
                device_id = item.data(Qt.UserRole)
                device = self.plugin.device_manager.get_device(device_id)
                if device and device not in selected_devices:
                    selected_devices.append(device)
        
        # Check if any devices are selected
        if not selected_devices:
            QMessageBox.warning(
                self,
                "No Devices Selected",
                "Please select one or more devices to include in the report."
            )
            return
            
        # Determine report format file extension
        format_index = self.report_format.currentIndex()
        if format_index == 0:
            file_ext = "txt"
            file_type = "Text Files (*.txt)"
        elif format_index == 1:
            file_ext = "html"
            file_type = "HTML Files (*.html)"
        elif format_index == 2:
            file_ext = "xlsx"
            file_type = "Excel Files (*.xlsx)"
        else:
            file_ext = "docx"
            file_type = "Word Files (*.docx)"
            
        # Get output file
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Report",
            f"{self.report_title.text().replace(' ', '_')}.{file_ext}",
            f"{file_type};;All Files (*.*)"
        )
        
        if not file_path:
            return
            
        # Validate date range if provided
        date_from = None
        date_to = None
        
        if self.date_from.text().strip():
            try:
                date_from = datetime.datetime.strptime(
                    self.date_from.text().strip(), 
                    "%Y-%m-%d"
                )
            except ValueError:
                QMessageBox.warning(
                    self,
                    "Invalid Date Format",
                    "Please enter dates in YYYY-MM-DD format."
                )
                return
                
        if self.date_to.text().strip():
            try:
                date_to = datetime.datetime.strptime(
                    self.date_to.text().strip(),
                    "%Y-%m-%d"
                )
                # Set to end of day
                date_to = date_to.replace(hour=23, minute=59, second=59)
            except ValueError:
                QMessageBox.warning(
                    self,
                    "Invalid Date Format",
                    "Please enter dates in YYYY-MM-DD format."
                )
                return
                
        # Generate the report
        try:
            if format_index == 0:
                # Text report
                self._generate_text_report(
                    file_path, 
                    selected_devices, 
                    self.report_title.text(),
                    self.include_device_info.isChecked(),
                    self.include_all_commands.isChecked(),
                    date_from,
                    date_to,
                    self.success_only.isChecked()
                )
            elif format_index == 1:
                # HTML report
                self._generate_html_report(
                    file_path, 
                    selected_devices, 
                    self.report_title.text(),
                    self.include_device_info.isChecked(),
                    self.include_all_commands.isChecked(),
                    date_from,
                    date_to,
                    self.success_only.isChecked()
                )
            elif format_index == 2:
                # Excel report
                self._generate_excel_report(
                    file_path, 
                    selected_devices, 
                    self.report_title.text(),
                    self.include_device_info.isChecked(),
                    self.include_all_commands.isChecked(),
                    date_from,
                    date_to,
                    self.success_only.isChecked()
                )
            else:
                # Word report
                self._generate_word_report(
                    file_path, 
                    selected_devices, 
                    self.report_title.text(),
                    self.include_device_info.isChecked(),
                    self.include_all_commands.isChecked(),
                    date_from,
                    date_to,
                    self.success_only.isChecked()
                )
                
            # Show success message
            QMessageBox.information(
                self,
                "Report Generated",
                f"Report saved to {file_path}"
            )
            
            # Close dialog
            self.accept()
            
        except Exception as e:
            # Show error message
            QMessageBox.critical(
                self,
                "Error Generating Report",
                f"An error occurred while generating the report: {str(e)}"
            )
            
    def _generate_text_report(self, file_path, devices, title, include_device_info, 
                             include_all, date_from, date_to, success_only):
        """Generate a text report"""
        with open(file_path, "w") as f:
            # Write report title
            f.write(f"{title}\n")
            f.write("=" * len(title) + "\n\n")
            
            # Write date
            current_date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            f.write(f"Generated: {current_date}\n\n")
            
            # Process each device
            for device in devices:
                # Write device header
                device_name = device.get_property("alias", "Unnamed Device")
                f.write(f"Device: {device_name}\n")
                f.write("-" * (len(device_name) + 8) + "\n\n")
                
                # Write device info if requested
                if include_device_info:
                    f.write("Device Information:\n")
                    for key, value in device.get_properties().items():
                        f.write(f"  {key}: {value}\n")
                    f.write("\n")
                
                # Get command outputs for this device
                outputs = self.plugin.get_command_outputs(device.id)
                
                if not outputs:
                    f.write("No command outputs available for this device.\n\n")
                    continue
                    
                # Process command outputs
                f.write("Command Outputs:\n")
                f.write("-" * 16 + "\n\n")
                
                # Track if we've written any outputs for this device
                wrote_outputs = False
                
                for cmd_id, cmd_outputs in outputs.items():
                    for timestamp, data in sorted(cmd_outputs.items()):
                        # Apply date filter if provided
                        if date_from or date_to:
                            dt = datetime.datetime.fromisoformat(timestamp)
                            
                            if date_from and dt < date_from:
                                continue
                                
                            if date_to and dt > date_to:
                                continue
                                
                        # Apply success filter if requested
                        if success_only and not data.get("success", True):
                            continue
                            
                        # Write command output
                        cmd_text = data.get("command", cmd_id)
                        output = data.get("output", "")
                        success = data.get("success", True)
                        
                        dt = datetime.datetime.fromisoformat(timestamp)
                        date_str = dt.strftime("%Y-%m-%d %H:%M:%S")
                        
                        f.write(f"Command: {cmd_text}\n")
                        f.write(f"Date/Time: {date_str}\n")
                        f.write(f"Success: {'Yes' if success else 'No'}\n")
                        f.write("Output:\n")
                        f.write("-" * 7 + "\n")
                        f.write(output + "\n\n")
                        
                        wrote_outputs = True
                        
                        # If we're not including all, just show the most recent
                        if not include_all:
                            break
                
                if not wrote_outputs:
                    f.write("No matching command outputs for this device.\n\n")
                    
                f.write("\n" + "=" * 50 + "\n\n")
                
    def _generate_html_report(self, file_path, devices, title, include_device_info, 
                             include_all, date_from, date_to, success_only):
        """Generate an HTML report"""
        with open(file_path, "w") as f:
            # Write HTML header
            f.write(f"""<!DOCTYPE html>
<html>
<head>
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        h1 {{ color: #2c3e50; }}
        h2 {{ color: #3498db; margin-top: 30px; }}
        h3 {{ color: #2980b9; }}
        pre {{ background-color: #f5f5f5; padding: 10px; border: 1px solid #ddd; overflow-x: auto; }}
        .device-info {{ background-color: #eef; padding: 10px; border: 1px solid #ddf; margin-bottom: 20px; }}
        .command {{ background-color: #efe; padding: 10px; border: 1px solid #dfd; margin-top: 20px; }}
        .command-failed {{ background-color: #fee; padding: 10px; border: 1px solid #fdd; margin-top: 20px; }}
        .timestamp {{ color: #777; font-style: italic; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p class="timestamp">Generated: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
""")
            
            # Process each device
            for device in devices:
                # Write device header
                device_name = device.get_property("alias", "Unnamed Device")
                f.write(f'    <h2>Device: {device_name}</h2>\n')
                
                # Write device info if requested
                if include_device_info:
                    f.write('    <div class="device-info">\n')
                    f.write('        <h3>Device Information</h3>\n')
                    f.write('        <table>\n')
                    
                    for key, value in device.get_properties().items():
                        f.write(f'            <tr><td><strong>{key}:</strong></td><td>{value}</td></tr>\n')
                        
                    f.write('        </table>\n')
                    f.write('    </div>\n')
                
                # Get command outputs for this device
                outputs = self.plugin.get_command_outputs(device.id)
                
                if not outputs:
                    f.write('    <p>No command outputs available for this device.</p>\n\n')
                    continue
                    
                # Process command outputs
                f.write('    <h3>Command Outputs</h3>\n')
                
                # Track if we've written any outputs for this device
                wrote_outputs = False
                
                for cmd_id, cmd_outputs in outputs.items():
                    for timestamp, data in sorted(cmd_outputs.items()):
                        # Apply date filter if provided
                        if date_from or date_to:
                            dt = datetime.datetime.fromisoformat(timestamp)
                            
                            if date_from and dt < date_from:
                                continue
                                
                            if date_to and dt > date_to:
                                continue
                                
                        # Apply success filter if requested
                        if success_only and not data.get("success", True):
                            continue
                            
                        # Write command output
                        cmd_text = data.get("command", cmd_id)
                        output = data.get("output", "")
                        success = data.get("success", True)
                        
                        dt = datetime.datetime.fromisoformat(timestamp)
                        date_str = dt.strftime("%Y-%m-%d %H:%M:%S")
                        
                        div_class = "command" if success else "command-failed"
                        
                        f.write(f'    <div class="{div_class}">\n')
                        f.write(f'        <h4>Command: {cmd_text}</h4>\n')
                        f.write(f'        <p>Date/Time: {date_str}</p>\n')
                        f.write(f'        <p>Success: {"Yes" if success else "No"}</p>\n')
                        f.write(f'        <h5>Output:</h5>\n')
                        f.write(f'        <pre>{output}</pre>\n')
                        f.write(f'    </div>\n')
                        
                        wrote_outputs = True
                        
                        # If we're not including all, just show the most recent
                        if not include_all:
                            break
                
                if not wrote_outputs:
                    f.write('    <p>No matching command outputs for this device.</p>\n\n')
                    
            # Write HTML footer
            f.write("""</body>
</html>
""")
                
    def _generate_excel_report(self, file_path, devices, title, include_device_info, 
                               include_all, date_from, date_to, success_only):
        """Generate an Excel report"""
        try:
            # Try to import openpyxl
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise RuntimeError("Excel report generation requires the openpyxl module. Please install it with 'pip install openpyxl'.")
            
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Command Report"
        
        # Set up styles
        title_font = Font(size=14, bold=True)
        header_font = Font(size=12, bold=True)
        device_font = Font(size=12, bold=True, color="0000FF")
        cmd_font = Font(size=11, bold=True)
        
        # Write title
        ws["A1"] = title
        ws["A1"].font = title_font
        ws.merge_cells("A1:G1")
        
        # Write generation timestamp
        ws["A2"] = f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        ws.merge_cells("A2:G2")
        
        row = 4
        
        # Process each device
        for device in devices:
            # Write device header
            device_name = device.get_property("alias", "Unnamed Device")
            ws[f"A{row}"] = f"Device: {device_name}"
            ws[f"A{row}"].font = device_font
            ws.merge_cells(f"A{row}:G{row}")
            row += 1
            
            # Write device info if requested
            if include_device_info:
                ws[f"A{row}"] = "Device Information:"
                ws[f"A{row}"].font = header_font
                ws.merge_cells(f"A{row}:G{row}")
                row += 1
                
                for key, value in device.get_properties().items():
                    ws[f"A{row}"] = key
                    ws[f"B{row}"] = str(value)
                    ws.merge_cells(f"B{row}:G{row}")
                    row += 1
                    
                row += 1
            
            # Get command outputs for this device
            outputs = self.plugin.get_command_outputs(device.id)
            
            if not outputs:
                ws[f"A{row}"] = "No command outputs available for this device."
                ws.merge_cells(f"A{row}:G{row}")
                row += 2
                continue
                
            # Process command outputs
            ws[f"A{row}"] = "Command Outputs:"
            ws[f"A{row}"].font = header_font
            ws.merge_cells(f"A{row}:G{row}")
            row += 1
            
            # Track if we've written any outputs for this device
            wrote_outputs = False
            
            for cmd_id, cmd_outputs in outputs.items():
                for timestamp, data in sorted(cmd_outputs.items()):
                    # Apply date filter if provided
                    if date_from or date_to:
                        dt = datetime.datetime.fromisoformat(timestamp)
                        
                        if date_from and dt < date_from:
                            continue
                            
                        if date_to and dt > date_to:
                            continue
                            
                    # Apply success filter if requested
                    if success_only and not data.get("success", True):
                        continue
                        
                    # Write command output
                    cmd_text = data.get("command", cmd_id)
                    output = data.get("output", "")
                    success = data.get("success", True)
                    
                    dt = datetime.datetime.fromisoformat(timestamp)
                    
                    # Command header
                    ws[f"A{row}"] = "Command:"
                    ws[f"B{row}"] = cmd_text
                    ws[f"B{row}"].font = cmd_font
                    ws.merge_cells(f"B{row}:G{row}")
                    row += 1
                    
                    # Date and success
                    ws[f"A{row}"] = "Date/Time:"
                    ws[f"B{row}"] = dt
                    
                    ws[f"C{row}"] = "Success:"
                    ws[f"D{row}"] = "Yes" if success else "No"
                    
                    row += 1
                    
                    # Output
                    ws[f"A{row}"] = "Output:"
                    ws.merge_cells(f"A{row}:G{row}")
                    row += 1
                    
                    # Split output into lines and write
                    output_lines = output.split("\n")
                    for line in output_lines:
                        ws[f"A{row}"] = line
                        ws.merge_cells(f"A{row}:G{row}")
                        row += 1
                        
                    row += 1
                    wrote_outputs = True
                    
                    # If we're not including all, just show the most recent
                    if not include_all:
                        break
            
            if not wrote_outputs:
                ws[f"A{row}"] = "No matching command outputs for this device."
                ws.merge_cells(f"A{row}:G{row}")
                row += 1
                
            row += 2
            
        # Adjust column widths
        for col in range(1, 8):
            ws.column_dimensions[get_column_letter(col)].width = 15
            
        # Save workbook
        wb.save(file_path)
        
    def _generate_word_report(self, file_path, devices, title, include_device_info, 
                              include_all, date_from, date_to, success_only):
        """Generate a Word report"""
        try:
            # Try to import python-docx
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("Word report generation requires the python-docx module. Please install it with 'pip install python-docx'.")
            
        # Create document
        doc = Document()
        
        # Set title
        title_para = doc.add_heading(title, level=0)
        
        # Add generation timestamp
        timestamp_para = doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Process each device
        for device in devices:
            # Add device header
            device_name = device.get_property("alias", "Unnamed Device")
            device_para = doc.add_heading(f"Device: {device_name}", level=1)
            
            # Add device info if requested
            if include_device_info:
                doc.add_heading("Device Information", level=2)
                
                # Create a table for device properties
                property_table = doc.add_table(rows=1, cols=2)
                property_table.style = "Table Grid"
                
                # Add header row
                header_cells = property_table.rows[0].cells
                header_cells[0].text = "Property"
                header_cells[1].text = "Value"
                
                # Add properties
                for key, value in device.get_properties().items():
                    row_cells = property_table.add_row().cells
                    row_cells[0].text = key
                    row_cells[1].text = str(value)
                    
                doc.add_paragraph()
            
            # Get command outputs for this device
            outputs = self.plugin.get_command_outputs(device.id)
            
            if not outputs:
                doc.add_paragraph("No command outputs available for this device.")
                doc.add_paragraph()
                continue
                
            # Process command outputs
            doc.add_heading("Command Outputs", level=2)
            
            # Track if we've written any outputs for this device
            wrote_outputs = False
            
            for cmd_id, cmd_outputs in outputs.items():
                for timestamp, data in sorted(cmd_outputs.items()):
                    # Apply date filter if provided
                    if date_from or date_to:
                        dt = datetime.datetime.fromisoformat(timestamp)
                        
                        if date_from and dt < date_from:
                            continue
                            
                        if date_to and dt > date_to:
                            continue
                            
                    # Apply success filter if requested
                    if success_only and not data.get("success", True):
                        continue
                        
                    # Write command output
                    cmd_text = data.get("command", cmd_id)
                    output = data.get("output", "")
                    success = data.get("success", True)
                    
                    dt = datetime.datetime.fromisoformat(timestamp)
                    date_str = dt.strftime("%Y-%m-%d %H:%M:%S")
                    
                    # Command header
                    cmd_heading = doc.add_heading(level=3)
                    cmd_heading.add_run(f"Command: {cmd_text}")
                    
                    # Date and success
                    info_para = doc.add_paragraph()
                    info_para.add_run(f"Date/Time: {date_str}\n")
                    info_para.add_run(f"Success: {'Yes' if success else 'No'}")
                    
                    # Output
                    doc.add_heading("Output", level=4)
                    output_para = doc.add_paragraph()
                    output_para.style = "No Spacing"
                    output_run = output_para.add_run(output)
                    output_run.font.name = "Courier New"
                    output_run.font.size = Pt(9)
                    
                    doc.add_paragraph()
                    wrote_outputs = True
                    
                    # If we're not including all, just show the most recent
                    if not include_all:
                        break
            
            if not wrote_outputs:
                doc.add_paragraph("No matching command outputs for this device.")
                
            # Add page break after each device except the last
            if device != devices[-1]:
                doc.add_page_break()
                
        # Save document
        doc.save(file_path) 