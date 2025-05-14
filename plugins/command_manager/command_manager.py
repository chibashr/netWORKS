#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Command Manager Plugin for NetWORKS
"""

import os
import json
import time
import uuid
import base64
import datetime
from pathlib import Path
from loguru import logger

from PySide6.QtCore import Qt, Signal, Slot, QObject
from PySide6.QtWidgets import (
    QMenu, QTabWidget, QWidget, QDialog, QVBoxLayout, 
    QHBoxLayout, QLabel, QPushButton, QTableWidget, QTableWidgetItem, 
    QHeaderView, QComboBox, QLineEdit, QTextEdit, QMessageBox,
    QFileDialog, QToolBar, QGroupBox, QFormLayout, QCheckBox,
    QDialogButtonBox, QSplitter, QTreeWidget, QTreeWidgetItem
)
from PySide6.QtGui import QIcon, QAction, QFont

# Import interfaces from the main application
from src.core.plugin_interface import PluginInterface

# Local imports
from plugins.command_manager.ui.command_dialog import CommandDialog
from plugins.command_manager.ui.credential_manager import CredentialManager
from plugins.command_manager.ui.command_output_panel import CommandOutputPanel
from plugins.command_manager.ui.command_set_editor import CommandSetEditor

# Import utilities
from plugins.command_manager.utils.encryption import encrypt_password, decrypt_password
from plugins.command_manager.utils.ssh_client import SSHClient
from plugins.command_manager.utils.telnet_client import TelnetClient
from plugins.command_manager.utils.command_set import CommandSet, Command
from plugins.command_manager.utils.credential_store import CredentialStore


class CommandManagerPlugin(PluginInterface):
    """Command Manager Plugin for NetWORKS"""
    
    def __init__(self):
        """Initialize the plugin"""
        super().__init__()
        
        logger.debug("Initializing Command Manager Plugin")
        
        # Plugin data directories
        self.data_dir = None
        self.commands_dir = None
        self.output_dir = None
        
        # UI components
        self.command_dialog = None
        self.output_panel = None
        self.toolbar_action = None
        self.toolbar = None
        self.context_menu_actions = {}
        
        # Data components
        self.command_sets = {}  # {device_type: {firmware: CommandSet}}
        self.credential_store = None
        self.outputs = {}       # {device_id: {command_id: {timestamp: output}}}
        
        logger.debug("Command Manager Plugin instance initialized")
        
    def initialize(self, app, plugin_info):
        """Initialize the plugin"""
        self.app = app
        self.plugin_info = plugin_info
        self.main_window = app.main_window
        self.device_manager = app.device_manager
        
        logger.debug(f"Command Manager initialization started. App: {app}, plugin_info: {plugin_info}")
        
        # Create UI components
        self._create_ui_components()
        
        # Create toolbar
        self.toolbar = self._create_toolbar()
        
        # Create output panel
        self.output_panel = CommandOutputPanel(self)
        
        # Create data directories
        self._create_data_directories()
        
        # Initialize outputs storage
        self.outputs = {}
        
        # Load command outputs
        self._load_command_outputs()
        
        # Create credential store
        try:
            data_dir = Path(self.plugin_info.path) / "data"
            self.credential_store = CredentialStore(data_dir)
            logger.debug(f"Credential store initialized with data_dir: {data_dir}")
        except Exception as e:
            logger.error(f"Error initializing credential store: {e}")
            logger.exception("Exception details:")
            self.credential_store = None
        
        # Load default command sets
        logger.debug("About to load default command sets")
        self._load_default_command_sets()
        logger.debug(f"Command sets loaded: {list(self.command_sets.keys())}")
        
        # Connect signals
        self._connect_signals()
        
        # Register context menu items directly with the device table
        logger.debug("Registering context menu actions with device table")
        try:
            # Check if device_table is available
            if hasattr(self.main_window, 'device_table'):
                # Register command manager action
                self.main_window.device_table.register_context_menu_action(
                    "Run Commands",
                    self._on_device_context_run_commands,
                    priority=550
                )
                
                # Register credential manager action
                self.main_window.device_table.register_context_menu_action(
                    "Manage Credentials",
                    self._on_device_context_credentials,
                    priority=560
                )
                
                logger.debug("Context menu actions registered successfully")
            else:
                logger.warning("device_table not found in main_window")
        except Exception as e:
            logger.error(f"Error registering context menu actions: {e}")
            logger.exception("Exception details:")
        
        logger.info("Command Manager plugin initialized successfully")
        return True
        
    def start(self):
        """Start the plugin"""
        logger.info(f"Starting {self.plugin_info.name} plugin")
        
        try:
            # Add toolbar to main window
            logger.debug(f"Adding toolbar to main window: {self.toolbar}")
            if self.toolbar is None:
                logger.error("Toolbar is None! Creating it now...")
                self.toolbar = self._create_toolbar()
            
            if self.main_window is None:
                logger.error("Main window is None! Cannot add toolbar.")
            else:
                try:
                    self.main_window.addToolBar(self.toolbar)
                    logger.debug("Toolbar added to main window")
                except Exception as e:
                    logger.error(f"Error adding toolbar to main window: {e}")
                    logger.exception("Exception details:")
            
            # Legacy device context menu items via device_manager (keeping for compatibility)
            logger.debug("Adding device context menu items via device_manager")
            try:
                if hasattr(self.device_manager, 'add_context_menu_item'):
                    self.device_manager.add_context_menu_item(
                        "Run Commands",
                        self._on_device_context_run_commands
                    )
                    
                    self.device_manager.add_context_menu_item(
                        "Manage Credentials",
                        self._on_device_context_credentials
                    )
                    logger.debug("Context menu items added via device_manager")
                else:
                    logger.debug("add_context_menu_item not found in device_manager")
            except Exception as e:
                logger.error(f"Error adding context menu items via device_manager: {e}")
                logger.exception("Exception details:")
            
            # Add tabs to device details
            logger.debug("Adding tabs to device details")
            try:
                self.device_manager.add_device_tab_provider(self)
                logger.debug("Tab provider added")
            except Exception as e:
                logger.error(f"Error adding tab provider: {e}")
                logger.exception("Exception details:")
            
            logger.info(f"{self.plugin_info.name} plugin started successfully")
            return True
        except Exception as e:
            logger.error(f"Error starting {self.plugin_info.name} plugin: {e}")
            logger.exception("Exception details:")
            return False
        
    def stop(self):
        """Stop the plugin"""
        if not super().stop():
            return False
            
        # Disconnect signals
        self._disconnect_signals()
        
        logger.info("Command Manager Plugin stopped")
        return True
        
    def cleanup(self):
        """Clean up plugin resources"""
        logger.info(f"{self.plugin_info.name} Plugin cleaned up")
        
        try:
            # Save command sets
            if hasattr(self, '_save_command_sets') and hasattr(self, 'command_sets') and self.command_sets:
                try:
                    self._save_command_sets()
                except Exception as e:
                    logger.error(f"Error saving command sets: {e}")
                    logger.exception("Exception details:")
            
            # Save command outputs
            if hasattr(self, '_save_command_outputs') and hasattr(self, 'outputs') and self.outputs:
                try:
                    self._save_command_outputs()
                except Exception as e:
                    logger.error(f"Error saving command outputs: {e}")
                    logger.exception("Exception details:")
                    
            # Close command dialog if open
            if hasattr(self, 'command_dialog') and self.command_dialog:
                try:
                    self.command_dialog.close()
                    self.command_dialog = None
                except Exception as e:
                    logger.error(f"Error closing command dialog: {e}")
            
            # Clean up UI
            if hasattr(self, 'output_panel'):
                self.output_panel = None
            
            # Unregister context menu actions
            if hasattr(self, 'main_window') and self.main_window and hasattr(self.main_window, 'device_table'):
                try:
                    # Unregister context menu actions
                    logger.debug("Unregistering context menu actions")
                    self.main_window.device_table.unregister_context_menu_action("Run Commands")
                    self.main_window.device_table.unregister_context_menu_action("Manage Credentials")
                    logger.debug("Context menu actions unregistered")
                except Exception as e:
                    logger.error(f"Error unregistering context menu actions: {e}")
                    logger.exception("Exception details:")
                    
            # Disconnect signals
            if hasattr(self, '_disconnect_signals'):
                try:
                    self._disconnect_signals()
                except Exception as e:
                    logger.error(f"Error disconnecting signals: {e}")
            
            return True
        except Exception as e:
            logger.error(f"Error during cleanup: {e}")
            logger.exception("Exception details:")
            return False
        
    def get_toolbar_actions(self):
        """Get actions to be added to the toolbar"""
        return [self.toolbar_action]
        
    def get_menu_actions(self):
        """Get actions to be added to the menu"""
        return {
            "Tools": [
                self.toolbar_action,
                self.credential_manager_action  # Add credential manager to Tools menu
            ]
        }
        
    def get_device_context_menu_actions(self):
        """Get actions to be added to the device context menu"""
        return list(self.context_menu_actions.values())
        
    def get_device_tabs(self, device):
        """Get tabs to be added to the device details view
        
        Args:
            device: The device object
            
        Returns:
            list: List of (tab_name, tab_widget) tuples
        """
        # Create a commands tab for this device
        commands_tab = self.create_device_command_tab(device)
        
        # Create a command output panel specifically for this device
        output_panel = CommandOutputPanel(self, device)
        
        # Return tabs with Commands tab first for better visibility
        return [
            ("Commands", commands_tab),
            ("Command Outputs", output_panel)
        ]

    def create_device_command_tab(self, device):
        """Create a tab to display device command history"""
        # Create widget
        widget = QWidget()
        layout = QVBoxLayout(widget)
        
        # Command output list
        device_command_list = QTableWidget()
        device_command_list.setColumnCount(4)
        device_command_list.setHorizontalHeaderLabels(["Command", "Date/Time", "Success", "View"])
        device_command_list.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        device_command_list.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        device_command_list.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
        device_command_list.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeToContents)
        device_command_list.setSelectionBehavior(QTableWidget.SelectRows)
        
        # Store device for reference
        device_command_list.setProperty("device", device)
        
        # Button bar
        button_layout = QHBoxLayout()
        
        refresh_btn = QPushButton("Refresh")
        refresh_btn.clicked.connect(lambda: self._refresh_device_commands(device_command_list))
        
        export_btn = QPushButton("Export")
        export_btn.clicked.connect(lambda: self._export_device_commands(device))
        
        delete_btn = QPushButton("Delete")
        delete_btn.clicked.connect(lambda: self._delete_device_commands(device_command_list))
        
        button_layout.addWidget(refresh_btn)
        button_layout.addWidget(export_btn)
        button_layout.addWidget(delete_btn)
        button_layout.addStretch()
        
        # Add widgets to layout
        layout.addWidget(device_command_list)
        layout.addLayout(button_layout)
        
        # Set up device command list
        self._refresh_device_commands(device_command_list)
        
        return widget
        
    def _refresh_device_commands(self, command_list):
        """Refresh the device command list"""
        # Clear the table
        command_list.setRowCount(0)
        
        # Get device
        device = command_list.property("device")
        if not device:
            return
            
        # Get command outputs for device
        outputs = self.get_command_outputs(device.id)
        
        # Add rows for each command output
        for cmd_id, cmd_outputs in outputs.items():
            for timestamp, data in cmd_outputs.items():
                row = command_list.rowCount()
                command_list.insertRow(row)
                
                # Command
                cmd_text = data.get("command", cmd_id)
                cmd_item = QTableWidgetItem(cmd_text)
                
                # Date/time
                dt = datetime.datetime.fromisoformat(timestamp)
                dt_item = QTableWidgetItem(dt.strftime("%Y-%m-%d %H:%M:%S"))
                
                # Success
                success = "Yes" if data.get("success", True) else "No"
                success_item = QTableWidgetItem(success)
                
                # View button
                view_btn = QPushButton("View")
                view_btn.setProperty("device_id", device.id)
                view_btn.setProperty("command_id", cmd_id)
                view_btn.setProperty("timestamp", timestamp)
                view_btn.clicked.connect(self._on_view_command)
                
                # Add to row
                command_list.setItem(row, 0, cmd_item)
                command_list.setItem(row, 1, dt_item)
                command_list.setItem(row, 2, success_item)
                command_list.setCellWidget(row, 3, view_btn)
    
    def _on_view_command(self):
        """Handle view command button"""
        # Get button that was clicked
        button = self.sender()
        if not button:
            return
            
        # Get properties
        device_id = button.property("device_id")
        command_id = button.property("command_id")
        timestamp = button.property("timestamp")
        
        # Get command output
        outputs = self.get_command_outputs(device_id, command_id)
        if not outputs or timestamp not in outputs:
            return
            
        output = outputs[timestamp]["output"]
        
        # Show output in dialog
        dialog = QDialog(self.main_window)
        dialog.setWindowTitle("Command Output")
        dialog.resize(600, 400)
        
        layout = QVBoxLayout(dialog)
        
        output_text = QTextEdit()
        output_text.setReadOnly(True)
        output_text.setFont(QFont("Courier New", 10))
        output_text.setText(output)
        
        layout.addWidget(output_text)
        
        buttons = QDialogButtonBox(QDialogButtonBox.Close)
        buttons.rejected.connect(dialog.reject)
        
        layout.addWidget(buttons)
        
        dialog.exec()
        
    def _export_device_commands(self, device):
        """Export device commands to file"""
        # Get selected device
        devices = self.device_manager.get_selected_devices()
        if not devices:
            return
            
        device = devices[0]
        
        # Ask for file to save to
        file_path, _ = QFileDialog.getSaveFileName(
            self.main_window,
            "Export Command Outputs",
            "",
            "Text Files (*.txt);;HTML Files (*.html);;All Files (*.*)"
        )
        
        if not file_path:
            return
            
        # Get command outputs for device
        outputs = self.get_command_outputs(device.id)
        
        try:
            # Determine export format
            if file_path.lower().endswith(".html"):
                # HTML export
                with open(file_path, "w") as f:
                    f.write("<html><head><title>Command Outputs</title></head><body>\n")
                    f.write(f"<h1>Command Outputs for {device.get_property('alias', 'Device')}</h1>\n")
                    
                    for cmd_id, cmd_outputs in outputs.items():
                        for timestamp, data in cmd_outputs.items():
                            dt = datetime.datetime.fromisoformat(timestamp)
                            cmd_text = data.get("command", cmd_id)
                            output = data.get("output", "")
                            
                            f.write(f"<h2>{cmd_text}</h2>\n")
                            f.write(f"<p>Date/Time: {dt.strftime('%Y-%m-%d %H:%M:%S')}</p>\n")
                            f.write("<pre>\n")
                            f.write(output)
                            f.write("\n</pre>\n")
                            f.write("<hr>\n")
                            
                    f.write("</body></html>")
            else:
                # Text export
                with open(file_path, "w") as f:
                    f.write(f"Command Outputs for {device.get_property('alias', 'Device')}\n")
                    f.write("=" * 50 + "\n\n")
                    
                    for cmd_id, cmd_outputs in outputs.items():
                        for timestamp, data in cmd_outputs.items():
                            dt = datetime.datetime.fromisoformat(timestamp)
                            cmd_text = data.get("command", cmd_id)
                            output = data.get("output", "")
                            
                            f.write(f"Command: {cmd_text}\n")
                            f.write(f"Date/Time: {dt.strftime('%Y-%m-%d %H:%M:%S')}\n")
                            f.write("-" * 50 + "\n")
                            f.write(output)
                            f.write("\n\n" + "=" * 50 + "\n\n")
            
            # Show success message
            QMessageBox.information(
                self.main_window,
                "Export Successful",
                f"Command outputs exported to {file_path}"
            )
            
        except Exception as e:
            # Show error message
            QMessageBox.critical(
                self.main_window,
                "Export Failed",
                f"Failed to export command outputs: {str(e)}"
            )
    
    def _delete_device_commands(self, command_list):
        """Delete selected device commands"""
        # Get selected rows
        selected_rows = command_list.selectedItems()
        if not selected_rows:
            return
            
        # Get selected device
        devices = self.device_manager.get_selected_devices()
        if not devices:
            return
            
        device = devices[0]
        
        # Confirm deletion
        result = QMessageBox.question(
            self.main_window,
            "Confirm Deletion",
            "Are you sure you want to delete the selected command outputs?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if result != QMessageBox.Yes:
            return
            
        # Delete selected outputs
        deleted = 0
        rows_to_delete = set()
        
        for item in selected_rows:
            row = item.row()
            rows_to_delete.add(row)
            
        for row in sorted(rows_to_delete, reverse=True):
            view_btn = command_list.cellWidget(row, 3)
            if view_btn:
                device_id = view_btn.property("device_id")
                command_id = view_btn.property("command_id")
                timestamp = view_btn.property("timestamp")
                
                if self.delete_command_output(device_id, command_id, timestamp):
                    deleted += 1
                    
        # Refresh the list
        self._refresh_device_commands(command_list)
        
        # Show success message
        QMessageBox.information(
            self.main_window,
            "Deletion Successful",
            f"Deleted {deleted} command output(s)"
        )
    
    def _create_data_directories(self):
        """Create data directories for the plugin"""
        # Main data directory
        self.data_dir = Path(self.plugin_info.path) / "data"
        self.data_dir.mkdir(parents=True, exist_ok=True)
        
        # Command sets directory
        self.commands_dir = self.data_dir / "commands"
        self.commands_dir.mkdir(exist_ok=True)
        
        # Command outputs directory
        self.output_dir = self.data_dir / "outputs"
        self.output_dir.mkdir(exist_ok=True)
        
    def _create_ui_components(self):
        """Create UI components for the plugin"""
        logger.debug("Creating UI components")
        
        # Create toolbar action (this is different from the toolbar itself)
        self.toolbar_action = QAction("Command Manager", self.main_window)
        self.toolbar_action.setToolTip("Manage and run commands on devices")
        self.toolbar_action.triggered.connect(self._show_command_dialog)
        logger.debug(f"Created toolbar_action: {self.toolbar_action}")
        
        # Create context menu actions
        self.context_menu_actions = {}
        
        run_action = QAction("Run Commands", self.main_window)
        run_action.triggered.connect(self._on_context_run_commands)
        self.context_menu_actions["run_commands"] = run_action
        
        # Enhance the credential management action
        manage_creds_action = QAction("Manage Device Credentials", self.main_window)
        manage_creds_action.setToolTip("Configure username, password and connection settings")
        manage_creds_action.triggered.connect(self._on_context_manage_credentials)
        manage_creds_action.setObjectName("ContextMenuCredentialAction")
        self.context_menu_actions["manage_credentials"] = manage_creds_action
        
        logger.debug(f"Created context menu actions: {list(self.context_menu_actions.keys())}")
        
        # Create output panel for device details
        self.output_panel = CommandOutputPanel(self)
        
        logger.debug("UI components created successfully")
        
    def _connect_signals(self):
        """Connect signals to slots"""
        logger.debug("Connecting signals")
        try:
            # Connect to device manager signals
            if hasattr(self.device_manager, 'device_added'):
                self.device_manager.device_added.connect(self._on_device_added)
            else:
                logger.warning("device_added signal not found")
                
            if hasattr(self.device_manager, 'device_removed'):
                self.device_manager.device_removed.connect(self._on_device_removed)
            else:
                logger.warning("device_removed signal not found")
                
            if hasattr(self.device_manager, 'device_changed'):
                self.device_manager.device_changed.connect(self._on_device_changed)
            else:
                logger.warning("device_changed signal not found")
                
            if hasattr(self.device_manager, 'selection_changed'):
                self.device_manager.selection_changed.connect(self._on_selection_changed)
            else:
                logger.warning("selection_changed signal not found")
                
            logger.debug("Signals connected successfully")
        except Exception as e:
            logger.error(f"Error connecting signals: {e}")
            logger.exception("Exception details:")
    
    def _disconnect_signals(self):
        """Disconnect signals from slots"""
        logger.debug("Disconnecting signals")
        try:
            # Only disconnect if the signal exists and we're connected
            # This prevents the "Failed to disconnect" warnings
            
            if hasattr(self.device_manager, 'device_added'):
                try:
                    self.device_manager.device_added.disconnect(self._on_device_added)
                    logger.debug("Disconnected device_added signal")
                except (RuntimeError, TypeError):
                    logger.debug("device_added signal was not connected")
                    
            if hasattr(self.device_manager, 'device_removed'):
                try:
                    self.device_manager.device_removed.disconnect(self._on_device_removed)
                    logger.debug("Disconnected device_removed signal")
                except (RuntimeError, TypeError):
                    logger.debug("device_removed signal was not connected")
                    
            if hasattr(self.device_manager, 'device_changed'):
                try:
                    self.device_manager.device_changed.disconnect(self._on_device_changed)
                    logger.debug("Disconnected device_changed signal")
                except (RuntimeError, TypeError):
                    logger.debug("device_changed signal was not connected")
                    
            if hasattr(self.device_manager, 'selection_changed'):
                try:
                    self.device_manager.selection_changed.disconnect(self._on_selection_changed)
                    logger.debug("Disconnected selection_changed signal")
                except (RuntimeError, TypeError):
                    logger.debug("selection_changed signal was not connected")
                    
            logger.debug("Signals disconnected successfully")
        except Exception as e:
            logger.error(f"Error disconnecting signals: {e}")
            logger.exception("Exception details:")
    
    # Event handlers for device manager signals
    
    def _on_device_added(self, device):
        """Handle device added event"""
        # Update UI if necessary
        if self.command_dialog:
            self.command_dialog.refresh_devices()
            
        if self.output_panel:
            self.output_panel.refresh()
    
    def _on_device_removed(self, device):
        """Handle device removed event"""
        # Update UI if necessary
        if self.command_dialog:
            self.command_dialog.refresh_devices()
            
        if self.output_panel:
            self.output_panel.refresh()
    
    def _on_device_changed(self, device):
        """Handle device changed event"""
        # Update UI if necessary
        if self.command_dialog:
            self.command_dialog.refresh_devices()
            
        if self.output_panel:
            self.output_panel.refresh()
    
    def _on_selection_changed(self, devices):
        """Handle selection changed event"""
        # Update the output panel with the selected device
        if self.output_panel and devices:
            self.output_panel.set_device(devices[0])
            
        # Update the commands panel if it exists
        if hasattr(self, 'commands_panel_widget') and self.commands_panel_widget and devices:
            self._update_commands_panel(devices[0])
            
    def _update_commands_panel(self, device):
        """Update the commands panel with the selected device
        
        Args:
            device: The selected device
        """
        logger.debug(f"Updating commands panel for device: {device.id}")
        
        # Clear existing content
        if hasattr(self, 'commands_panel_widget') and self.commands_panel_widget:
            # Remove all widgets from the layout
            while self.commands_panel_widget.layout().count():
                item = self.commands_panel_widget.layout().takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            
            # Create a splitter to divide command list and output viewer
            splitter = QSplitter(Qt.Vertical)
            
            # Top section - Command list
            top_widget = QWidget()
            top_layout = QVBoxLayout(top_widget)
            top_layout.setContentsMargins(5, 5, 5, 5)
            
            # Command output list
            device_command_list = QTableWidget()
            device_command_list.setColumnCount(3)  # Removed the View column
            device_command_list.setHorizontalHeaderLabels(["Command", "Date/Time", "Success"])
            device_command_list.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
            device_command_list.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
            device_command_list.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
            device_command_list.setSelectionBehavior(QTableWidget.SelectRows)
            device_command_list.setSelectionMode(QTableWidget.SingleSelection)
            
            # Store device for reference
            device_command_list.setProperty("device", device)
            
            top_layout.addWidget(device_command_list)
            
            # Bottom section - Command output viewer
            bottom_widget = QWidget()
            bottom_layout = QVBoxLayout(bottom_widget)
            bottom_layout.setContentsMargins(5, 5, 5, 5)
            
            # Output viewer header with view toggle
            header_layout = QHBoxLayout()
            
            # Output viewer label
            viewer_label = QLabel("Command Output:")
            viewer_label.setStyleSheet("font-weight: bold;")
            header_layout.addWidget(viewer_label)
            
            header_layout.addStretch()
            
            # View toggle buttons
            view_toggle_group = QWidget()
            view_toggle_layout = QHBoxLayout(view_toggle_group)
            view_toggle_layout.setContentsMargins(0, 0, 0, 0)
            view_toggle_layout.setSpacing(0)
            
            raw_btn = QPushButton("Raw")
            raw_btn.setCheckable(True)
            raw_btn.setChecked(True)
            raw_btn.setMaximumWidth(60)
            
            table_btn = QPushButton("Table")
            table_btn.setCheckable(True)
            table_btn.setMaximumWidth(60)
            
            # Style the toggle buttons as a segmented control
            raw_btn.setStyleSheet("""
                QPushButton {
                    border: 1px solid #8f8f8f;
                    border-top-left-radius: 4px;
                    border-bottom-left-radius: 4px;
                    border-right: none;
                    padding: 4px 8px;
                    background-color: #f0f0f0;
                }
                QPushButton:checked {
                    background-color: #d0d0d0;
                }
            """)
            
            table_btn.setStyleSheet("""
                QPushButton {
                    border: 1px solid #8f8f8f;
                    border-top-right-radius: 4px;
                    border-bottom-right-radius: 4px;
                    padding: 4px 8px;
                    background-color: #f0f0f0;
                }
                QPushButton:checked {
                    background-color: #d0d0d0;
                }
            """)
            
            view_toggle_layout.addWidget(raw_btn)
            view_toggle_layout.addWidget(table_btn)
            
            header_layout.addWidget(view_toggle_group)
            
            bottom_layout.addLayout(header_layout)
            
            # Create stacked widget for different views
            output_stack = QTabWidget()
            output_stack.setDocumentMode(True)
            output_stack.tabBar().setVisible(False)
            
            # Raw text view
            output_text = QTextEdit()
            output_text.setReadOnly(True)
            output_text.setFont(QFont("Courier New", 10))
            output_text.setPlaceholderText("Select a command to view its output")
            output_stack.addTab(output_text, "Raw")
            
            # Table view
            output_table = QTableWidget()
            output_table.setEditTriggers(QTableWidget.NoEditTriggers)
            output_stack.addTab(output_table, "Table")
            
            # Store references for later use
            device_command_list.setProperty("output_text", output_text)
            device_command_list.setProperty("output_table", output_table)
            device_command_list.setProperty("output_stack", output_stack)
            device_command_list.setProperty("raw_btn", raw_btn)
            device_command_list.setProperty("table_btn", table_btn)
            
            bottom_layout.addWidget(output_stack)
            
            # Connect toggle buttons
            def toggle_raw():
                if raw_btn.isChecked():
                    table_btn.setChecked(False)
                    output_stack.setCurrentIndex(0)
                elif not table_btn.isChecked():
                    # If neither is checked, check raw
                    raw_btn.setChecked(True)
                    output_stack.setCurrentIndex(0)
            
            def toggle_table():
                if table_btn.isChecked():
                    raw_btn.setChecked(False)
                    output_stack.setCurrentIndex(1)
                elif not raw_btn.isChecked():
                    # If neither is checked, check table
                    table_btn.setChecked(True)
                    output_stack.setCurrentIndex(1)
            
            raw_btn.clicked.connect(toggle_raw)
            table_btn.clicked.connect(toggle_table)
            
            # Connect selection change to update the viewer
            device_command_list.itemSelectionChanged.connect(
                lambda: self._on_command_selection_changed(device_command_list)
            )
            
            # Add widgets to splitter
            splitter.addWidget(top_widget)
            splitter.addWidget(bottom_widget)
            splitter.setStretchFactor(0, 1)
            splitter.setStretchFactor(1, 2)
            
            # Button bar for both top and bottom sections
            button_layout = QHBoxLayout()
            
            # Command action buttons
            run_btn = QPushButton("Run Command")
            run_btn.setToolTip("Run a new command on this device")
            run_btn.clicked.connect(lambda: self._run_command_for_device(device))
            
            refresh_btn = QPushButton("Refresh")
            refresh_btn.setToolTip("Refresh the command list")
            refresh_btn.clicked.connect(lambda: self._refresh_device_commands(device_command_list))
            
            export_btn = QPushButton("Export")
            export_btn.setToolTip("Export command outputs to file")
            export_btn.clicked.connect(lambda: self._export_device_commands(device))
            
            # Control buttons for selected command
            view_btn = QPushButton("View")
            view_btn.setToolTip("View selected command in a popup window")
            view_btn.setEnabled(False)
            view_btn.clicked.connect(lambda: self._view_selected_command(device_command_list))
            
            delete_btn = QPushButton("Delete")
            delete_btn.setToolTip("Delete selected command output")
            delete_btn.setEnabled(False)
            delete_btn.clicked.connect(lambda: self._delete_selected_command(device_command_list))
            
            # Store references to enable/disable these buttons
            device_command_list.setProperty("view_button", view_btn)
            device_command_list.setProperty("delete_button", delete_btn)
            
            # Add buttons to layout
            button_layout.addWidget(run_btn)
            button_layout.addWidget(refresh_btn)
            button_layout.addWidget(export_btn)
            button_layout.addWidget(view_btn)
            button_layout.addWidget(delete_btn)
            button_layout.addStretch()
            
            # Add widgets to main layout
            self.commands_panel_widget.layout().addWidget(splitter)
            self.commands_panel_widget.layout().addLayout(button_layout)
            
            # Load commands for this device
            self._refresh_device_commands(device_command_list)
            
    def _on_command_selection_changed(self, command_list):
        """Handle command selection changed
        
        Args:
            command_list: The command list widget
        """
        selected_items = command_list.selectedItems()
        
        # Get output widgets
        output_text = command_list.property("output_text")
        output_table = command_list.property("output_table")
        output_stack = command_list.property("output_stack")
        raw_btn = command_list.property("raw_btn")
        table_btn = command_list.property("table_btn")
        
        # Enable/disable buttons based on selection
        view_btn = command_list.property("view_button")
        delete_btn = command_list.property("delete_button")
        
        if selected_items and view_btn and delete_btn:
            view_btn.setEnabled(True)
            delete_btn.setEnabled(True)
            
            # Get the selected row
            row = selected_items[0].row()
            
            # Get device, command and timestamp
            device = command_list.property("device")
            cmd_text = command_list.item(row, 0).text()
            timestamp = command_list.item(row, 1).data(Qt.UserRole)
            cmd_id = command_list.item(row, 0).data(Qt.UserRole)
            
            # Get command output
            outputs = self.get_command_outputs(device.id, cmd_id)
            
            if outputs and timestamp in outputs:
                output = outputs[timestamp]["output"]
                success = outputs[timestamp]["success"]
                
                # Update the raw text output
                if output_text:
                    output_text.setText(output)
                
                # Try to parse as table and update table view
                if output_table:
                    has_table = self._update_table_view(output, output_table)
                    
                    # Enable/disable table view based on whether we could parse a table
                    if table_btn:
                        table_btn.setEnabled(has_table)
                        
                        # If table view was selected but we can't show a table, switch to raw
                        if output_stack and output_stack.currentIndex() == 1 and not has_table:
                            output_stack.setCurrentIndex(0)
                            if raw_btn:
                                raw_btn.setChecked(True)
                            if table_btn:
                                table_btn.setChecked(False)
            else:
                # Clear outputs
                if output_text:
                    output_text.setText(f"Output not found for command: {cmd_text}")
                if output_table:
                    output_table.setRowCount(0)
                    output_table.setColumnCount(0)
                if table_btn:
                    table_btn.setEnabled(False)
        else:
            # No selection, disable buttons
            if view_btn:
                view_btn.setEnabled(False)
            if delete_btn:
                delete_btn.setEnabled(False)
                
            # Clear outputs
            if output_text:
                output_text.clear()
            if output_table:
                output_table.setRowCount(0)
                output_table.setColumnCount(0)
            if table_btn:
                table_btn.setEnabled(False)
                
    def _update_table_view(self, output, table_widget):
        """Try to parse command output as a table and update the table widget
        
        Args:
            output: Command output text
            table_widget: Table widget to update
            
        Returns:
            bool: True if output could be parsed as a table, False otherwise
        """
        # Clear the table
        table_widget.setRowCount(0)
        table_widget.setColumnCount(0)
        
        if not output:
            return False
            
        # Split into lines
        lines = output.strip().split('\n')
        if len(lines) < 2:
            return False
            
        # Try to detect table structure
        # Look for lines that might be headers (contain multiple consecutive spaces)
        header_candidates = []
        data_lines = []
        
        for i, line in enumerate(lines):
            # Skip empty lines
            if not line.strip():
                continue
                
            # Look for header-like patterns
            if '  ' in line and i < len(lines) / 2:  # Headers are usually near the top
                header_candidates.append((i, line))
            elif line.strip():
                data_lines.append((i, line))
                
        # No headers found
        if not header_candidates:
            return False
            
        # Use the last header candidate
        header_index, header_line = header_candidates[-1]
        
        # Try to extract column positions from the header
        col_positions = self._find_column_positions(header_line)
        
        # If we couldn't find at least 2 columns, abort
        if len(col_positions) < 2:
            return False
            
        # Extract column names
        col_names = []
        for i in range(len(col_positions)):
            start = col_positions[i]
            end = col_positions[i+1] if i < len(col_positions) - 1 else len(header_line)
            col_name = header_line[start:end].strip()
            col_names.append(col_name)
            
        # Set up the table
        table_widget.setColumnCount(len(col_names))
        table_widget.setHorizontalHeaderLabels(col_names)
        
        # Populate with data
        row_index = 0
        for i, line in data_lines:
            # Skip lines before or equal to the header
            if i <= header_index:
                continue
                
            # Skip separator lines or empty lines
            if not line.strip() or all(c in '-+' for c in line.strip()):
                continue
                
            # Extract columns based on positions
            cols = []
            for j in range(len(col_positions)):
                start = col_positions[j]
                end = col_positions[j+1] if j < len(col_positions) - 1 else len(line)
                
                # Handle if line is shorter than header
                if start >= len(line):
                    cols.append("")
                else:
                    cols.append(line[start:end].strip())
            
            # Add to table
            table_widget.insertRow(row_index)
            for col_index, col_value in enumerate(cols):
                item = QTableWidgetItem(col_value)
                table_widget.setItem(row_index, col_index, item)
            row_index += 1
            
        # If we added rows, adjust column widths
        if table_widget.rowCount() > 0:
            table_widget.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
            for i in range(1, table_widget.columnCount()):
                table_widget.horizontalHeader().setSectionResizeMode(i, QHeaderView.ResizeToContents)
            return True
            
        return False
        
    def _find_column_positions(self, line):
        """Find starting positions of columns in a header line
        
        Args:
            line: Header line text
            
        Returns:
            list: List of column start positions
        """
        positions = [0]  # Always start at position 0
        in_whitespace = False
        
        for i, c in enumerate(line):
            if c.isspace():
                in_whitespace = True
            elif in_whitespace:
                # Transition from whitespace to non-whitespace
                in_whitespace = False
                # Check if we have at least 2 spaces before this position
                if i >= 2 and line[i-2:i].isspace():
                    positions.append(i)
                    
        return positions
    
    def _view_selected_command(self, command_list):
        """View the selected command in a popup window
        
        Args:
            command_list: The command list widget
        """
        selected_items = command_list.selectedItems()
        if not selected_items:
            return
            
        # Get the selected row
        row = selected_items[0].row()
        
        # Get device, command and timestamp
        device = command_list.property("device")
        cmd_text = command_list.item(row, 0).text()
        timestamp = command_list.item(row, 1).data(Qt.UserRole)
        cmd_id = command_list.item(row, 0).data(Qt.UserRole)
        
        # Get command output
        outputs = self.get_command_outputs(device.id, cmd_id)
        
        if outputs and timestamp in outputs:
            output = outputs[timestamp]["output"]
            
            # Show output in dialog
            dialog = QDialog(self.main_window)
            dialog.setWindowTitle("Command Output")
            dialog.resize(800, 600)
            
            layout = QVBoxLayout(dialog)
            
            # Add command info
            info_label = QLabel(f"Command: {cmd_text}")
            info_label.setWordWrap(True)
            layout.addWidget(info_label)
            
            # Add timestamp
            dt = datetime.datetime.fromisoformat(timestamp)
            date_str = dt.strftime("%Y-%m-%d %H:%M:%S")
            time_label = QLabel(f"Date/Time: {date_str}")
            layout.addWidget(time_label)
            
            # Add tab widget for raw/table views
            tab_widget = QTabWidget()
            
            # Raw output tab
            raw_tab = QWidget()
            raw_layout = QVBoxLayout(raw_tab)
            
            output_text = QTextEdit()
            output_text.setReadOnly(True)
            output_text.setFont(QFont("Courier New", 10))
            output_text.setText(output)
            raw_layout.addWidget(output_text)
            
            tab_widget.addTab(raw_tab, "Raw Output")
            
            # Table output tab
            table_tab = QWidget()
            table_layout = QVBoxLayout(table_tab)
            
            output_table = QTableWidget()
            output_table.setEditTriggers(QTableWidget.NoEditTriggers)
            table_layout.addWidget(output_table)
            
            # Try to parse as table
            has_table = self._update_table_view(output, output_table)
            
            # Only add table tab if we could parse it
            if has_table:
                tab_widget.addTab(table_tab, "Table View")
            
            layout.addWidget(tab_widget)
            
            # Add buttons
            buttons = QDialogButtonBox(QDialogButtonBox.Close)
            buttons.rejected.connect(dialog.reject)
            layout.addWidget(buttons)
            
            dialog.exec()
            
    def _delete_selected_command(self, command_list):
        """Delete the selected command output
        
        Args:
            command_list: The command list widget
        """
        selected_items = command_list.selectedItems()
        if not selected_items:
            return
            
        # Confirm deletion
        result = QMessageBox.question(
            self.main_window,
            "Confirm Deletion",
            "Are you sure you want to delete this command output?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if result != QMessageBox.Yes:
            return
            
        # Get the selected row
        row = selected_items[0].row()
        
        # Get device, command and timestamp
        device = command_list.property("device")
        cmd_id = command_list.item(row, 0).data(Qt.UserRole)
        timestamp = command_list.item(row, 1).data(Qt.UserRole)
        
        # Delete the command output
        if self.delete_command_output(device.id, cmd_id, timestamp):
            # Refresh the command list
            self._refresh_device_commands(command_list)
            
            # Show success message
            QMessageBox.information(
                self.main_window,
                "Deletion Successful",
                "Command output deleted successfully"
            )
        else:
            # Show error message
            QMessageBox.critical(
                self.main_window,
                "Deletion Failed",
                "Failed to delete command output"
            )
            
    def _run_command_for_device(self, device):
        """Run a command for the selected device
        
        Args:
            device: The device to run a command on
        """
        # Create a dialog to select and run a command
        dialog = CommandDialog(self, [device], self.main_window)
        dialog.exec()
            
    def _refresh_device_commands(self, command_list):
        """Refresh the device command list"""
        # Get output widgets
        output_text = command_list.property("output_text")
        output_table = command_list.property("output_table")
        table_btn = command_list.property("table_btn")
        
        # Clear the table
        command_list.setRowCount(0)
        
        # Clear output text if provided
        if output_text:
            output_text.clear()
            
        # Clear table view if provided
        if output_table:
            output_table.setRowCount(0)
            output_table.setColumnCount(0)
            
        # Disable table button
        if table_btn:
            table_btn.setEnabled(False)
            
        # Clear button states
        view_btn = command_list.property("view_button")
        delete_btn = command_list.property("delete_button")
        
        if view_btn:
            view_btn.setEnabled(False)
            
        if delete_btn:
            delete_btn.setEnabled(False)
        
        # Get device
        device = command_list.property("device")
        if not device:
            return
            
        # Get command outputs for device
        outputs = self.get_command_outputs(device.id)
        
        # Add rows for each command output
        for cmd_id, cmd_outputs in outputs.items():
            for timestamp, data in cmd_outputs.items():
                row = command_list.rowCount()
                command_list.insertRow(row)
                
                # Command
                cmd_text = data.get("command", cmd_id)
                cmd_item = QTableWidgetItem(cmd_text)
                cmd_item.setData(Qt.UserRole, cmd_id)
                
                # Date/time
                dt = datetime.datetime.fromisoformat(timestamp)
                dt_item = QTableWidgetItem(dt.strftime("%Y-%m-%d %H:%M:%S"))
                dt_item.setData(Qt.UserRole, timestamp)
                
                # Success
                success = "Yes" if data.get("success", True) else "No"
                success_item = QTableWidgetItem(success)
                
                # Add to row
                command_list.setItem(row, 0, cmd_item)
                command_list.setItem(row, 1, dt_item)
                command_list.setItem(row, 2, success_item)
    
    # UI Action handlers
    
    def _show_command_dialog(self):
        """Show the command dialog"""
        if not self.command_dialog:
            self.command_dialog = CommandDialog(self, parent=self.main_window)
            
        # Show dialog (non-modal)
        self.command_dialog.show()
        self.command_dialog.raise_()
        self.command_dialog.activateWindow()
    
    def _on_context_run_commands(self):
        """Handle run commands context menu action"""
        # Get selected devices
        devices = self.device_manager.get_selected_devices()
        
        if not devices:
            QMessageBox.warning(
                self.main_window,
                "No Devices Selected",
                "Please select one or more devices to run commands on."
            )
            return
        
        # Show command dialog with selected devices
        self._show_command_dialog()
        self.command_dialog.set_selected_devices(devices)
    
    def _on_context_manage_credentials(self):
        """Handle manage credentials context menu action"""
        # Get selected devices
        devices = self.device_manager.get_selected_devices()
        
        if not devices:
            QMessageBox.warning(
                self.main_window,
                "No Devices Selected",
                "Please select one or more devices to manage credentials for."
            )
            return
        
        # Show credential manager dialog
        cred_manager = CredentialManager(self, devices, self.main_window)
        cred_manager.exec()
    
    def _on_run_commands(self):
        """Handle run commands action"""
        # Get selected devices
        devices = self.device_manager.get_selected_devices()
        if not devices:
            QMessageBox.warning(
                self.main_window,
                "No Devices Selected",
                "Please select one or more devices to run commands on."
            )
            return
            
        # Open command dialog
        dialog = CommandDialog(self, devices, parent=self.main_window)
        dialog.exec()

    def _on_manage_sets(self):
        """Handle manage command sets action"""
        dialog = CommandSetEditor(self)
        dialog.exec()

    def _create_toolbar(self):
        """Create plugin toolbar"""
        logger.debug("Creating Command Manager toolbar")
        self.toolbar = QToolBar("Command Manager")
        self.toolbar.setObjectName("CommandManagerToolbar")  # Setting object name helps avoid issues
        
        try:
            # Action to run commands
            run_action = QAction("Run Commands", self.main_window)
            run_action.setToolTip("Run commands on devices")
            run_action.triggered.connect(self._on_run_commands)
            logger.debug(f"Created run_action: {run_action}")
            self.toolbar.addAction(run_action)
            
            # Action to manage command sets
            sets_action = QAction("Command Sets", self.main_window)
            sets_action.setToolTip("Manage command sets")
            sets_action.triggered.connect(self._on_manage_sets)
            logger.debug(f"Created sets_action: {sets_action}")
            self.toolbar.addAction(sets_action)
            
            # Add a separator before the credential management button to make it stand out
            self.toolbar.addSeparator()
            
            # Action to manage credentials - make it more prominent with bold styling
            creds_action = QAction("👤 Credential Manager", self.main_window)
            creds_action.setToolTip("Manage device credentials and access settings")
            creds_action.triggered.connect(self._on_manage_credentials)
            # Set a unique object name to help with debugging
            creds_action.setObjectName("CredentialManagerAction")
            
            # Store the action as an instance variable so we can reference it later
            self.credential_manager_action = creds_action
            
            logger.debug(f"Created creds_action: {creds_action}")
            self.toolbar.addAction(creds_action)
            
            # Add a separator after the credential management button
            self.toolbar.addSeparator()
            
            # Action to generate reports
            report_action = QAction("Generate Report", self.main_window)
            report_action.setToolTip("Generate command output reports")
            report_action.triggered.connect(self._on_generate_report)
            logger.debug(f"Created report_action: {report_action}")
            self.toolbar.addAction(report_action)
            
        except Exception as e:
            logger.error(f"Error creating toolbar actions: {e}")
            logger.exception("Exception details:")
        
        # Ensure the toolbar has some actions
        if not self.toolbar.actions():
            logger.warning("Toolbar has no actions - applying fallback action")
            fallback = QAction("Command Manager", self.main_window)
            fallback.setToolTip("Command Manager")
            fallback.triggered.connect(lambda: logger.info("Fallback action triggered"))
            self.toolbar.addAction(fallback)
        
        logger.debug(f"Toolbar created with {len(self.toolbar.actions())} actions")
        return self.toolbar

    def _on_generate_report(self):
        """Handle generate report action"""
        # Create report generation dialog
        dialog = ReportGenerator(self, self.main_window)
        dialog.exec()

    def _on_manage_credentials(self):
        """Handle manage credentials action"""
        logger.info("Opening Credential Manager")
        
        # Create and display the credential manager dialog
        try:
            dialog = CredentialManager(self)
            logger.debug("Credential Manager dialog created successfully")
            
            # Set window title to be more descriptive
            dialog.setWindowTitle("Device Credential Manager")
            
            # Make the dialog a bit larger for better usability
            dialog.resize(700, 550)
            
            # Execute the dialog
            result = dialog.exec()
            
            if result == QDialog.Accepted:
                logger.info("Credential changes saved")
            else:
                logger.info("Credential Manager closed without saving")
                
        except Exception as e:
            logger.error(f"Error opening Credential Manager: {e}")
            logger.exception("Exception details:")
            QMessageBox.critical(
                self.main_window,
                "Error Opening Credential Manager",
                f"An error occurred while opening the Credential Manager: {str(e)}"
            )

    def _on_device_context_run_commands(self, devices):
        """Handle run commands context menu item"""
        dialog = CommandDialog(self, devices, parent=self.main_window)
        dialog.exec()
        
    def _on_device_context_credentials(self, devices):
        """Handle manage credentials context menu item"""
        logger.info(f"Opening Credential Manager for {len(devices)} selected devices")
        
        try:
            dialog = CredentialManager(self, devices, self.main_window)
            logger.debug("Credential Manager dialog created successfully")
            
            # Set window title to be more descriptive
            dialog.setWindowTitle("Device Credential Manager")
            
            # Make the dialog a bit larger for better usability
            dialog.resize(700, 550)
            
            # Execute the dialog
            result = dialog.exec()
            
            if result == QDialog.Accepted:
                logger.info("Credential changes saved")
            else:
                logger.info("Credential Manager closed without saving")
                
        except Exception as e:
            logger.error(f"Error opening Credential Manager: {e}")
            logger.exception("Exception details:")
            QMessageBox.critical(
                self.main_window,
                "Error Opening Credential Manager",
                f"An error occurred while opening the Credential Manager: {str(e)}"
            )

    def _load_default_command_sets(self):
        """Load default command sets"""
        logger.debug("Loading default command sets")
        
        # Check if cisco_iosxe.json exists in the commands directory
        cisco_iosxe_path = self.commands_dir / "cisco_iosxe.json"
        if not cisco_iosxe_path.exists():
            logger.debug(f"Default cisco_iosxe.json not found at {cisco_iosxe_path}")
            # Look for it in the plugin directory
            source_path = Path(self.plugin_info.path) / "data" / "commands" / "cisco_iosxe.json"
            logger.debug(f"Checking for source file at {source_path}")
            if source_path.exists():
                # Copy to commands directory
                try:
                    logger.debug(f"Source file found, copying to {cisco_iosxe_path}")
                    with open(source_path, "r") as src:
                        data = json.load(src)
                    
                    with open(cisco_iosxe_path, "w") as dst:
                        json.dump(data, dst, indent=2)
                        
                    logger.info(f"Copied default command set: cisco_iosxe.json")
                except Exception as e:
                    logger.error(f"Error copying default command set: {e}")
            else:
                logger.debug("Source file not found, will try to use embedded data")
        else:
            logger.debug(f"Default cisco_iosxe.json already exists at {cisco_iosxe_path}")
        
        # Load all command sets
        logger.debug("Calling _load_command_sets")
        self._load_command_sets()

    def _load_command_sets(self):
        """Load command sets from disk"""
        logger.debug("Loading command sets from disk")
        self.command_sets = {}
        
        # Check if the command sets directory exists
        if not self.commands_dir.exists():
            logger.warning(f"Command sets directory does not exist: {self.commands_dir}")
            self.commands_dir.mkdir(parents=True, exist_ok=True)
            return
        
        # Keep track of problematic files to potentially clean up
        problem_files = []
        
        # Iterate through command set files
        command_files = list(self.commands_dir.glob("*.json"))
        logger.debug(f"Found {len(command_files)} command set files: {[f.name for f in command_files]}")
        
        for file_path in command_files:
            try:
                # Skip empty files
                if file_path.stat().st_size == 0:
                    logger.warning(f"Skipping empty command set file: {file_path}")
                    problem_files.append(file_path)
                    continue
                    
                logger.info(f"Attempting to load command set from {file_path}")
                with open(file_path, "r", encoding='utf-8') as f:
                    data = json.load(f)
                
                # Extract device type and firmware from filename if it's a list of commands
                # This handles legacy files that are simply lists of commands
                if isinstance(data, list):
                    logger.debug(f"File {file_path.name} contains a list of commands. Extracting info from filename...")
                    
                    # Extract device type and firmware from filename
                    filename = file_path.stem
                    parts = filename.split('_')
                    
                    if len(parts) >= 2:
                        # Reconstruct full device type and firmware
                        device_parts = []
                        firmware_parts = []
                        
                        # Check if filename follows the pattern cisco_ios_xe_16_x
                        if "cisco" in parts[0].lower() and "ios" in filename.lower():
                            device_type = "Cisco IOS XE"
                            
                            # Get firmware from the last parts
                            if "16" in filename or "17" in filename:
                                # Find version in filename
                                for part in parts:
                                    if part.isdigit() or part.startswith(("16", "17")):
                                        firmware_parts.append(part)
                                
                                firmware = '.'.join(firmware_parts) if firmware_parts else "16.x"
                            else:
                                firmware = "16.x"  # Default if none found
                        else:
                            # Generic approach for other devices
                            device_type = ' '.join(parts[:-1]).title()
                            firmware = parts[-1].replace('_', '.')
                        
                        # Create a properly formatted command set
                        logger.debug(f"Extracted device_type={device_type}, firmware={firmware}")
                        commands = data
                    else:
                        # Can't determine device type and firmware from filename
                        device_type = "Unknown"
                        firmware = "Unknown"
                        commands = data
                        logger.warning(f"Could not determine device type and firmware from filename: {filename}")
                
                # Handle standard command set format
                elif isinstance(data, dict) and "device_type" in data and "firmware_version" in data and "commands" in data:
                    device_type = data["device_type"]
                    firmware = data["firmware_version"]
                    commands = data["commands"]
                else:
                    logger.warning(f"Invalid command set format in file: {file_path}")
                    logger.debug(f"Data structure: {type(data)}, Fields: {list(data.keys()) if isinstance(data, dict) else 'not a dict'}")
                    problem_files.append(file_path)
                    continue
                
                logger.debug(f"Loading command set for {device_type} ({firmware}) with {len(commands)} commands")
                
                if device_type not in self.command_sets:
                    self.command_sets[device_type] = {}
                    
                self.command_sets[device_type][firmware] = commands
                
                logger.info(f"Loaded command set: {device_type} ({firmware})")
                
            except json.JSONDecodeError as e:
                logger.error(f"Error parsing JSON in command set file {file_path}: {e}")
                problem_files.append(file_path)
                # Try to fix the file if it's the Cisco IOS XE one
                if "cisco" in file_path.name.lower():
                    logger.debug(f"Attempting to fix Cisco command set: {file_path}")
                    self._fix_cisco_command_set()
            except Exception as e:
                logger.error(f"Error loading command set from {file_path}: {e}")
                logger.exception("Exception details:")
                problem_files.append(file_path)
        
        # Clean up problematic files if we have at least one good command set
        if self.command_sets and problem_files:
            logger.debug(f"Found {len(problem_files)} problematic files to clean up")
            for path in problem_files:
                try:
                    # Don't delete cisco_iosxe.json as it might have been fixed
                    if path.name.lower() != "cisco_iosxe.json":
                        logger.debug(f"Removing problematic file: {path}")
                        path.unlink()
                        logger.info(f"Removed problematic command set file: {path}")
                except Exception as e:
                    logger.warning(f"Could not remove problematic file {path}: {e}")
                    
        # Ensure we have at least the default command set
        if not self.command_sets:
            logger.info("No command sets loaded, loading defaults")
            self._fix_cisco_command_set()
            # Try loading again after fixing
            self._load_command_sets()
            
        logger.debug(f"Final command sets: {list(self.command_sets.keys())}")

    def get_device_types(self):
        """Get available device types from command sets
        
        Returns:
            list: List of device types
        """
        logger.debug("Getting available device types")
        if not self.command_sets:
            logger.warning("No command sets loaded")
            return []
            
        device_types = list(self.command_sets.keys())
        logger.debug(f"Found {len(device_types)} device types: {device_types}")
        return device_types
        
    def get_firmware_versions(self, device_type):
        """Get available firmware versions for a device type
        
        Args:
            device_type (str): Device type to get firmware versions for
            
        Returns:
            list: List of firmware versions
        """
        logger.debug(f"Getting firmware versions for {device_type}")
        if not self.command_sets or device_type not in self.command_sets:
            logger.warning(f"Device type {device_type} not found in command sets")
            return []
            
        firmware_versions = list(self.command_sets[device_type].keys())
        logger.debug(f"Found {len(firmware_versions)} firmware versions: {firmware_versions}")
        return firmware_versions
        
    def get_commands(self, device_type, firmware_version):
        """Get commands for a device type and firmware version
        
        Args:
            device_type (str): Device type to get commands for
            firmware_version (str): Firmware version to get commands for
            
        Returns:
            list: List of commands
        """
        logger.debug(f"Getting commands for {device_type} ({firmware_version})")
        if (not self.command_sets or 
            device_type not in self.command_sets or 
            firmware_version not in self.command_sets[device_type]):
            logger.warning(f"Command set for {device_type} ({firmware_version}) not found")
            return []
            
        commands = self.command_sets[device_type][firmware_version]
        logger.debug(f"Found {len(commands)} commands")
        return commands

    def get_command_set(self, device_type, firmware_version):
        """Get a command set for a device type and firmware version
        
        This returns a CommandSet object for compatibility with CommandDialog
        
        Args:
            device_type (str): Device type to get commands for
            firmware_version (str): Firmware version to get commands for
            
        Returns:
            CommandSet: CommandSet object with commands
        """
        logger.debug(f"Getting command set for {device_type} ({firmware_version})")
        commands = self.get_commands(device_type, firmware_version)
        
        # Create a CommandSet object from the commands list
        from plugins.command_manager.utils.command_set import CommandSet, Command
        
        # Convert commands to Command objects if needed
        command_objects = []
        for cmd in commands:
            if not hasattr(cmd, 'alias'):
                # Convert dict to Command object
                if isinstance(cmd, dict):
                    alias = cmd.get('alias', cmd.get('command', 'Command'))
                    description = cmd.get('description', '')
                    command = cmd.get('command', '')
                    cmd_obj = Command(command, alias, description)
                    command_objects.append(cmd_obj)
                else:
                    # Skip invalid commands
                    logger.warning(f"Invalid command format: {cmd}")
            else:
                # Already a Command object
                command_objects.append(cmd)
                
        # Create CommandSet
        command_set = CommandSet(device_type, firmware_version, command_objects)
        return command_set

    def _save_command_sets(self):
        """Save command sets to disk"""
        logger.debug("Saving command sets to disk")
        # Check if the commands directory exists
        if not self.commands_dir.exists():
            self.commands_dir.mkdir(parents=True, exist_ok=True)
        
        # Iterate through command sets
        for device_type, firmware_sets in self.command_sets.items():
            for firmware, command_set in firmware_sets.items():
                try:
                    # Create filename from device_type and firmware
                    # Use lowercase to maintain consistent filenames
                    device_type_safe = device_type.lower().replace(' ', '_')
                    firmware_safe = firmware.replace('.', '_')
                    filename = f"{device_type_safe}_{firmware_safe}.json"
                    file_path = self.commands_dir / filename
                    
                    # Save to file - handle both CommandSet objects and plain lists/dicts
                    with open(file_path, "w") as f:
                        if hasattr(command_set, 'to_dict'):
                            # CommandSet object
                            json.dump(command_set.to_dict(), f, indent=2)
                        else:
                            # If it's a list, wrap it in the proper format
                            data_to_save = command_set
                            if isinstance(command_set, list):
                                data_to_save = {
                                    "device_type": device_type,
                                    "firmware_version": firmware,
                                    "commands": command_set
                                }
                            json.dump(data_to_save, f, indent=2)
                
                except Exception as e:
                    logger.error(f"Error saving command set {device_type}_{firmware}: {e}")
                    logger.exception("Exception details:")
                    
    def _save_command_outputs(self):
        """Save command outputs to disk"""
        logger.debug("Saving command outputs to disk")
        # Check if the outputs directory exists
        if not self.output_dir.exists():
            self.output_dir.mkdir(parents=True, exist_ok=True)
            
        # No outputs to save
        if not self.outputs:
            logger.debug("No command outputs to save")
            return
            
        # Save outputs for each device in its own folder in the plugin directory
        for device_id, commands in self.outputs.items():
            try:
                # Create device output directory in plugin folder
                device_dir = self.output_dir / device_id
                device_dir.mkdir(exist_ok=True)
                
                # Save the outputs to disk
                output_file = device_dir / "command_outputs.json"
                with open(output_file, "w") as f:
                    json.dump(commands, f, indent=2)
                    
                logger.debug(f"Saved command outputs for device {device_id} to {output_file}")
                
                # Also save to workspace device folder
                try:
                    # Get workspace device directory
                    workspace_device_dir = Path("config/workspaces/default/devices") / device_id
                    if workspace_device_dir.exists():
                        # Create commands directory if it doesn't exist
                        commands_dir = workspace_device_dir / "commands"
                        commands_dir.mkdir(exist_ok=True)
                        
                        # Save command outputs to workspace
                        workspace_output_file = commands_dir / "command_outputs.json"
                        with open(workspace_output_file, "w") as f:
                            json.dump(commands, f, indent=2)
                            
                        logger.debug(f"Saved command outputs to workspace: {workspace_output_file}")
                except Exception as e:
                    logger.error(f"Error saving command outputs to workspace for device {device_id}: {e}")
                    logger.exception("Exception details:")
                
            except Exception as e:
                logger.error(f"Error saving command outputs for device {device_id}: {e}")
                logger.exception("Exception details:")
        
        # Also save the full outputs file for backward compatibility
        try:
            # Save the outputs to disk
            output_file = self.output_dir / "command_outputs.json"
            with open(output_file, "w") as f:
                json.dump(self.outputs, f, indent=2)
                
            logger.debug(f"Saved command outputs to {output_file}")
        except Exception as e:
            logger.error(f"Error saving command outputs: {e}")
            logger.exception("Exception details:")

    def add_command_set(self, command_set):
        """Add or update a command set
        
        Args:
            command_set: CommandSet object to add or update
        """
        logger.debug(f"Adding command set: {command_set.device_type} ({command_set.firmware_version})")
        
        # Create device type entry if it doesn't exist
        if command_set.device_type not in self.command_sets:
            self.command_sets[command_set.device_type] = {}
            
        # Add or update command set
        self.command_sets[command_set.device_type][command_set.firmware_version] = command_set.commands
        
        # Save command sets
        self._save_command_sets()
        
        logger.info(f"Added command set: {command_set.device_type} ({command_set.firmware_version})")

    def get_device_credentials(self, device_id):
        """Get credentials for a device
        
        Args:
            device_id (str): Device ID to get credentials for
            
        Returns:
            dict: Credentials for the device
        """
        logger.debug(f"Getting credentials for device: {device_id}")
        
        # Check if we have a credential store
        if not hasattr(self, 'credential_store') or not self.credential_store:
            logger.warning("No credential store available")
            return {}
            
        # Get credentials from store - we need to use get_device_credentials not get_credentials
        credentials = self.credential_store.get_device_credentials(device_id)
        
        # If no credentials found, return empty dict
        if not credentials:
            logger.warning(f"No credentials found for device: {device_id}")
            return {}
            
        logger.debug(f"Got credentials for device: {device_id}")
        return credentials
        
    def run_command(self, device, command, credentials=None):
        """Run a command on a device
        
        Args:
            device: Device to run command on
            command (str): Command to run
            credentials (dict, optional): Credentials to use
            
        Returns:
            dict: Command result
        """
        logger.debug(f"Running command on device: {device.id}")
        
        # If no credentials provided, get them
        if not credentials:
            credentials = self.get_device_credentials(device.id)
            
        # Default result
        result = {
            "success": False,
            "output": f"Command: {command}\n\nNo connection method available for device: {device.id}"
        }
        
        # Get device properties
        device_type = device.get_property("device_type", "")
        ip_address = device.get_property("ip_address", "")
        
        # Log credential status
        if not credentials or not credentials.get("username"):
            logger.warning(f"No valid credentials found for device {device.id} ({ip_address})")
            result["output"] = f"Command: {command}\n\nNo valid credentials available for {ip_address}"
            return result
            
        # Determine connection type
        connection_type = credentials.get("connection_type", "ssh").lower()
        
        # Handle SSH connections
        if connection_type == "ssh":
            try:
                # Import SSH client
                from plugins.command_manager.utils.ssh_client import SSHClient
                
                logger.debug(f"Connecting to {ip_address} via SSH with username: {credentials.get('username')}")
                
                # Create SSH client with required parameters
                ssh = SSHClient(
                    host=ip_address,
                    username=credentials.get("username"),
                    password=credentials.get("password", ""),
                    enable_password=credentials.get("enable_password", "")
                )
                
                try:
                    # Connect to device
                    ssh.connect()
                    
                    # Try to enter enable mode if needed
                    if credentials.get("enable_password"):
                        ssh.enable()
                    
                    # Execute the command
                    output = ssh.execute(command)
                    
                    # Create successful result
                    result = {
                        "success": True,
                        "output": output
                    }
                    
                    # Disconnect cleanly
                    ssh.disconnect()
                    
                    logger.debug(f"SSH command execution completed successfully")
                except Exception as e:
                    logger.error(f"SSH execution error: {e}")
                    result["output"] = f"Command: {command}\n\nSSH Connection error: {str(e)}"
            except Exception as e:
                logger.error(f"Error running SSH command: {e}")
                logger.exception("Exception details:")
                result["output"] = f"Command: {command}\n\nError running SSH command: {str(e)}"
                
        # Handle Telnet connections
        elif connection_type == "telnet":
            try:
                # Import Telnet client
                from plugins.command_manager.utils.telnet_client import TelnetClient
                
                logger.debug(f"Connecting to {ip_address} via Telnet with username: {credentials.get('username')}")
                
                # Create Telnet client with required parameters
                telnet = TelnetClient(
                    host=ip_address,
                    username=credentials.get("username"),
                    password=credentials.get("password", ""),
                    enable_password=credentials.get("enable_password", "")
                )
                
                try:
                    # Connect to device
                    telnet.connect()
                    
                    # Try to enter enable mode if needed
                    if hasattr(telnet, 'enable') and credentials.get("enable_password"):
                        telnet.enable()
                    
                    # Execute the command
                    output = telnet.execute(command)
                    
                    # Create successful result
                    result = {
                        "success": True,
                        "output": output
                    }
                    
                    # Disconnect cleanly
                    telnet.disconnect()
                    
                    logger.debug(f"Telnet command execution completed successfully")
                except Exception as e:
                    logger.error(f"Telnet execution error: {e}")
                    result["output"] = f"Command: {command}\n\nTelnet Connection error: {str(e)}"
            except Exception as e:
                logger.error(f"Error running Telnet command: {e}")
                logger.exception("Exception details:")
                result["output"] = f"Command: {command}\n\nError running Telnet command: {str(e)}"
        else:
            logger.warning(f"Unsupported connection type: {connection_type}")
            result["output"] = f"Command: {command}\n\nUnsupported connection type: {connection_type}"
            
        return result
        
    def add_command_output(self, device_id, command_id, output, command_text=None):
        """Add command output to history
        
        Args:
            device_id (str): Device ID
            command_id (str): Command ID
            output (str): Command output
            command_text (str, optional): Command text
        """
        logger.debug(f"Adding command output for device: {device_id}, command: {command_id}")
        
        # Create device entry if it doesn't exist
        if device_id not in self.outputs:
            self.outputs[device_id] = {}
            
        # Create command entry if it doesn't exist
        if command_id not in self.outputs[device_id]:
            self.outputs[device_id][command_id] = {}
            
        # Add output with timestamp
        timestamp = datetime.datetime.now().isoformat()
        self.outputs[device_id][command_id][timestamp] = {
            "output": output,
            "success": True,
            "command": command_text if command_text else command_id
        }
        
        # Save outputs
        self._save_command_outputs()
        
        logger.debug(f"Added command output for device: {device_id}, command: {command_id}")
        
    def get_command_outputs(self, device_id, command_id=None):
        """Get command outputs for a device
        
        Args:
            device_id (str): Device ID to get outputs for
            command_id (str, optional): Command ID to get outputs for
            
        Returns:
            dict: Command outputs
        """
        logger.debug(f"Getting command outputs for device: {device_id}")
        
        # Check if we have outputs for this device
        if device_id not in self.outputs:
            logger.debug(f"No outputs found for device: {device_id}")
            return {}
            
        # If command_id is provided, get only those outputs
        if command_id:
            if command_id in self.outputs[device_id]:
                return self.outputs[device_id][command_id]
            else:
                logger.debug(f"No outputs found for device: {device_id}, command: {command_id}")
                return {}
                
        # Otherwise return all outputs for the device
        return self.outputs[device_id]

    def get_all_device_credentials(self):
        """Get all device credentials
        
        Returns:
            dict: Dictionary of device credentials by device ID
        """
        logger.debug("Getting all device credentials")
        
        # Check if we have a credential store
        if not hasattr(self, 'credential_store') or not self.credential_store:
            logger.warning("No credential store available")
            return {}
            
        # Get all device credentials
        try:
            return self.credential_store.get_all_device_credentials()
        except Exception as e:
            logger.error(f"Error getting all device credentials: {e}")
            logger.exception("Exception details:")
            return {}
            
    def get_all_group_credentials(self):
        """Get all group credentials
        
        Returns:
            dict: Dictionary of group credentials by group name
        """
        logger.debug("Getting all group credentials")
        
        # Check if we have a credential store
        if not hasattr(self, 'credential_store') or not self.credential_store:
            logger.warning("No credential store available")
            return {}
            
        # Get all group credentials
        try:
            return self.credential_store.get_all_group_credentials()
        except Exception as e:
            logger.error(f"Error getting all group credentials: {e}")
            logger.exception("Exception details:")
            return {}
            
    def get_all_subnet_credentials(self):
        """Get all subnet credentials
        
        Returns:
            dict: Dictionary of subnet credentials by subnet
        """
        logger.debug("Getting all subnet credentials")
        
        # Check if we have a credential store
        if not hasattr(self, 'credential_store') or not self.credential_store:
            logger.warning("No credential store available")
            return {}
            
        # Get all subnet credentials
        try:
            return self.credential_store.get_all_subnet_credentials()
        except Exception as e:
            logger.error(f"Error getting all subnet credentials: {e}")
            logger.exception("Exception details:")
            return {}
            
    def set_device_credentials(self, device_id, credentials):
        """Set credentials for a device
        
        Args:
            device_id (str): Device ID to set credentials for
            credentials (dict): Credentials to set
        """
        logger.debug(f"Setting credentials for device: {device_id}")
        
        # Check if we have a credential store
        if not hasattr(self, 'credential_store') or not self.credential_store:
            logger.warning("No credential store available")
            return
            
        # Set credentials in store
        try:
            self.credential_store.set_device_credentials(device_id, credentials)
            logger.info(f"Credentials set for device: {device_id}")
        except Exception as e:
            logger.error(f"Error setting credentials for device {device_id}: {e}")
            logger.exception("Exception details:")
            
    def set_group_credentials(self, group_name, credentials):
        """Set credentials for a device group
        
        Args:
            group_name (str): Group name to set credentials for
            credentials (dict): Credentials to set
        """
        logger.debug(f"Setting credentials for group: {group_name}")
        
        # Check if we have a credential store
        if not hasattr(self, 'credential_store') or not self.credential_store:
            logger.warning("No credential store available")
            return
            
        # Set credentials in store
        try:
            self.credential_store.set_group_credentials(group_name, credentials)
            logger.info(f"Credentials set for group: {group_name}")
        except Exception as e:
            logger.error(f"Error setting credentials for group {group_name}: {e}")
            logger.exception("Exception details:")
            
    def set_subnet_credentials(self, subnet, credentials):
        """Set credentials for a subnet
        
        Args:
            subnet (str): Subnet to set credentials for
            credentials (dict): Credentials to set
        """
        logger.debug(f"Setting credentials for subnet: {subnet}")
        
        # Check if we have a credential store
        if not hasattr(self, 'credential_store') or not self.credential_store:
            logger.warning("No credential store available")
            return
            
        # Set credentials in store
        try:
            self.credential_store.set_subnet_credentials(subnet, credentials)
            logger.info(f"Credentials set for subnet: {subnet}")
        except Exception as e:
            logger.error(f"Error setting credentials for subnet {subnet}: {e}")
            logger.exception("Exception details:")
            
    def delete_device_credentials(self, device_id):
        """Delete credentials for a device
        
        Args:
            device_id (str): Device ID to delete credentials for
        """
        logger.debug(f"Deleting credentials for device: {device_id}")
        
        # Check if we have a credential store
        if not hasattr(self, 'credential_store') or not self.credential_store:
            logger.warning("No credential store available")
            return
            
        # Delete credentials from store
        try:
            self.credential_store.delete_device_credentials(device_id)
            logger.info(f"Credentials deleted for device: {device_id}")
        except Exception as e:
            logger.error(f"Error deleting credentials for device {device_id}: {e}")
            logger.exception("Exception details:")
            
    def delete_group_credentials(self, group_name):
        """Delete credentials for a device group
        
        Args:
            group_name (str): Group name to delete credentials for
        """
        logger.debug(f"Deleting credentials for group: {group_name}")
        
        # Check if we have a credential store
        if not hasattr(self, 'credential_store') or not self.credential_store:
            logger.warning("No credential store available")
            return
            
        # Delete credentials from store
        try:
            self.credential_store.delete_group_credentials(group_name)
            logger.info(f"Credentials deleted for group: {group_name}")
        except Exception as e:
            logger.error(f"Error deleting credentials for group {group_name}: {e}")
            logger.exception("Exception details:")
            
    def delete_subnet_credentials(self, subnet):
        """Delete credentials for a subnet
        
        Args:
            subnet (str): Subnet to delete credentials for
        """
        logger.debug(f"Deleting credentials for subnet: {subnet}")
        
        # Check if we have a credential store
        if not hasattr(self, 'credential_store') or not self.credential_store:
            logger.warning("No credential store available")
            return
            
        # Delete credentials from store
        try:
            self.credential_store.delete_subnet_credentials(subnet)
            logger.info(f"Credentials deleted for subnet: {subnet}")
        except Exception as e:
            logger.error(f"Error deleting credentials for subnet {subnet}: {e}")
            logger.exception("Exception details:")

    def delete_command_output(self, device_id, command_id, timestamp=None):
        """Delete a command output from history
        
        Args:
            device_id (str): Device ID
            command_id (str): Command ID
            timestamp (str, optional): Specific timestamp to delete, or None to delete all
            
        Returns:
            bool: True if deleted successfully, False otherwise
        """
        logger.debug(f"Deleting command output for device: {device_id}, command: {command_id}, timestamp: {timestamp}")
        
        # Check if we have outputs for this device
        if device_id not in self.outputs:
            logger.warning(f"No outputs found for device: {device_id}")
            return False
            
        # Check if we have outputs for this command
        if command_id not in self.outputs[device_id]:
            logger.warning(f"No outputs found for device: {device_id}, command: {command_id}")
            return False
            
        # If timestamp is provided, delete only that timestamp
        if timestamp:
            if timestamp in self.outputs[device_id][command_id]:
                del self.outputs[device_id][command_id][timestamp]
                
                # If no more timestamps for this command, delete the command
                if not self.outputs[device_id][command_id]:
                    del self.outputs[device_id][command_id]
                    
                # If no more commands for this device, delete the device
                if not self.outputs[device_id]:
                    del self.outputs[device_id]
                    
                # Save outputs
                self._save_command_outputs()
                
                logger.debug(f"Deleted command output for device: {device_id}, command: {command_id}, timestamp: {timestamp}")
                return True
            else:
                logger.warning(f"No output found for device: {device_id}, command: {command_id}, timestamp: {timestamp}")
                return False
        # If no timestamp, delete all outputs for this command
        else:
            del self.outputs[device_id][command_id]
            
            # If no more commands for this device, delete the device
            if not self.outputs[device_id]:
                del self.outputs[device_id]
                
            # Save outputs
            self._save_command_outputs()
            
            logger.debug(f"Deleted all command outputs for device: {device_id}, command: {command_id}")
            return True

    def _load_command_outputs(self):
        """Load command outputs from disk"""
        logger.debug("Loading command outputs from disk")
        
        # Check if the outputs directory exists
        if not self.output_dir.exists():
            logger.debug(f"Output directory does not exist: {self.output_dir}")
            # Create it if not exists
            self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize outputs
        self.outputs = {}
        
        # First, try to load device-specific output files from plugin data directory
        device_dirs = [d for d in self.output_dir.iterdir() if d.is_dir()]
        for device_dir in device_dirs:
            device_id = device_dir.name
            output_file = device_dir / "command_outputs.json"
            
            if output_file.exists():
                try:
                    with open(output_file, "r") as f:
                        device_outputs = json.load(f)
                        self.outputs[device_id] = device_outputs
                        
                    logger.debug(f"Loaded command outputs for device {device_id} from {output_file}")
                except Exception as e:
                    logger.error(f"Error loading command outputs for device {device_id}: {e}")
                    logger.exception("Exception details:")
        
        # Next, try to load from workspace device folders
        workspace_device_dir = Path("config/workspaces/default/devices")
        if workspace_device_dir.exists():
            for device_dir in workspace_device_dir.iterdir():
                if device_dir.is_dir():
                    device_id = device_dir.name
                    commands_dir = device_dir / "commands"
                    if commands_dir.exists():
                        output_file = commands_dir / "command_outputs.json"
                        if output_file.exists():
                            try:
                                with open(output_file, "r") as f:
                                    device_outputs = json.load(f)
                                    # Only load if we don't already have outputs for this device
                                    if device_id not in self.outputs:
                                        self.outputs[device_id] = device_outputs
                                        logger.debug(f"Loaded command outputs for device {device_id} from workspace: {output_file}")
                            except Exception as e:
                                logger.error(f"Error loading command outputs for device {device_id} from workspace: {e}")
                                logger.exception("Exception details:")
            
        # If no device-specific outputs found, try to load from the legacy file
        if not self.outputs:
            legacy_file = self.output_dir / "command_outputs.json"
            if legacy_file.exists():
                try:
                    with open(legacy_file, "r") as f:
                        self.outputs = json.load(f)
                    
                    logger.debug(f"Loaded command outputs from legacy file {legacy_file}")
                    
                    # Migrate to new format
                    self._save_command_outputs()
                except Exception as e:
                    logger.error(f"Error loading legacy command outputs: {e}")
                    logger.exception("Exception details:")
        
            # Log some statistics
            device_count = len(self.outputs)
            command_count = 0
            output_count = 0
            
            for device_id, commands in self.outputs.items():
                command_count += len(commands)
                for cmd_id, timestamps in commands.items():
                    output_count += len(timestamps)
                    
            logger.info(f"Loaded {output_count} command outputs for {command_count} commands across {device_count} devices")

    def get_device_panels(self):
        """Get panels to be added to the device properties panel
        
        Returns:
            list: List of (panel_name, panel_widget) tuples
        """
        logger.debug("Getting device panels for Command Manager plugin")
        
        # Create a device-agnostic commands tab initially - it will update when a device is selected
        commands_tab = QWidget()
        layout = QVBoxLayout(commands_tab)
        
        # Add info label
        info_label = QLabel("Select a device to view and manage commands for that device.")
        info_label.setWordWrap(True)
        layout.addWidget(info_label)
        
        # We'll store this widget to update it later when a device is selected
        self.commands_panel_widget = commands_tab
        
        return [
            ("Commands", commands_tab)
        ]
        
    def _on_selection_changed(self, devices):
        """Handle selection changed event"""
        # Update the output panel with the selected device
        if self.output_panel and devices:
            self.output_panel.set_device(devices[0])
            
        # Update the commands panel if it exists
        if hasattr(self, 'commands_panel_widget') and self.commands_panel_widget and devices:
            self._update_commands_panel(devices[0])
            
    def _update_commands_panel(self, device):
        """Update the commands panel with the selected device
        
        Args:
            device: The selected device
        """
        logger.debug(f"Updating commands panel for device: {device.id}")
        
        # Clear existing content
        if hasattr(self, 'commands_panel_widget') and self.commands_panel_widget:
            # Remove all widgets from the layout
            while self.commands_panel_widget.layout().count():
                item = self.commands_panel_widget.layout().takeAt(0)
                if item.widget():
                    item.widget().deleteLater()
            
            # Create a splitter to divide command list and output viewer
            splitter = QSplitter(Qt.Vertical)
            
            # Top section - Command list
            top_widget = QWidget()
            top_layout = QVBoxLayout(top_widget)
            top_layout.setContentsMargins(5, 5, 5, 5)
            
            # Command output list
            device_command_list = QTableWidget()
            device_command_list.setColumnCount(3)  # Removed the View column
            device_command_list.setHorizontalHeaderLabels(["Command", "Date/Time", "Success"])
            device_command_list.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
            device_command_list.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
            device_command_list.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeToContents)
            device_command_list.setSelectionBehavior(QTableWidget.SelectRows)
            device_command_list.setSelectionMode(QTableWidget.SingleSelection)
            
            # Store device for reference
            device_command_list.setProperty("device", device)
            
            top_layout.addWidget(device_command_list)
            
            # Bottom section - Command output viewer
            bottom_widget = QWidget()
            bottom_layout = QVBoxLayout(bottom_widget)
            bottom_layout.setContentsMargins(5, 5, 5, 5)
            
            # Output viewer header with view toggle
            header_layout = QHBoxLayout()
            
            # Output viewer label
            viewer_label = QLabel("Command Output:")
            viewer_label.setStyleSheet("font-weight: bold;")
            header_layout.addWidget(viewer_label)
            
            header_layout.addStretch()
            
            # View toggle buttons
            view_toggle_group = QWidget()
            view_toggle_layout = QHBoxLayout(view_toggle_group)
            view_toggle_layout.setContentsMargins(0, 0, 0, 0)
            view_toggle_layout.setSpacing(0)
            
            raw_btn = QPushButton("Raw")
            raw_btn.setCheckable(True)
            raw_btn.setChecked(True)
            raw_btn.setMaximumWidth(60)
            
            table_btn = QPushButton("Table")
            table_btn.setCheckable(True)
            table_btn.setMaximumWidth(60)
            
            # Style the toggle buttons as a segmented control
            raw_btn.setStyleSheet("""
                QPushButton {
                    border: 1px solid #8f8f8f;
                    border-top-left-radius: 4px;
                    border-bottom-left-radius: 4px;
                    border-right: none;
                    padding: 4px 8px;
                    background-color: #f0f0f0;
                }
                QPushButton:checked {
                    background-color: #d0d0d0;
                }
            """)
            
            table_btn.setStyleSheet("""
                QPushButton {
                    border: 1px solid #8f8f8f;
                    border-top-right-radius: 4px;
                    border-bottom-right-radius: 4px;
                    padding: 4px 8px;
                    background-color: #f0f0f0;
                }
                QPushButton:checked {
                    background-color: #d0d0d0;
                }
            """)
            
            view_toggle_layout.addWidget(raw_btn)
            view_toggle_layout.addWidget(table_btn)
            
            header_layout.addWidget(view_toggle_group)
            
            bottom_layout.addLayout(header_layout)
            
            # Create stacked widget for different views
            output_stack = QTabWidget()
            output_stack.setDocumentMode(True)
            output_stack.tabBar().setVisible(False)
            
            # Raw text view
            output_text = QTextEdit()
            output_text.setReadOnly(True)
            output_text.setFont(QFont("Courier New", 10))
            output_text.setPlaceholderText("Select a command to view its output")
            output_stack.addTab(output_text, "Raw")
            
            # Table view
            output_table = QTableWidget()
            output_table.setEditTriggers(QTableWidget.NoEditTriggers)
            output_stack.addTab(output_table, "Table")
            
            # Store references for later use
            device_command_list.setProperty("output_text", output_text)
            device_command_list.setProperty("output_table", output_table)
            device_command_list.setProperty("output_stack", output_stack)
            device_command_list.setProperty("raw_btn", raw_btn)
            device_command_list.setProperty("table_btn", table_btn)
            
            bottom_layout.addWidget(output_stack)
            
            # Connect toggle buttons
            def toggle_raw():
                if raw_btn.isChecked():
                    table_btn.setChecked(False)
                    output_stack.setCurrentIndex(0)
                elif not table_btn.isChecked():
                    # If neither is checked, check raw
                    raw_btn.setChecked(True)
                    output_stack.setCurrentIndex(0)
            
            def toggle_table():
                if table_btn.isChecked():
                    raw_btn.setChecked(False)
                    output_stack.setCurrentIndex(1)
                elif not raw_btn.isChecked():
                    # If neither is checked, check table
                    table_btn.setChecked(True)
                    output_stack.setCurrentIndex(1)
            
            raw_btn.clicked.connect(toggle_raw)
            table_btn.clicked.connect(toggle_table)
            
            # Connect selection change to update the viewer
            device_command_list.itemSelectionChanged.connect(
                lambda: self._on_command_selection_changed(device_command_list)
            )
            
            # Add widgets to splitter
            splitter.addWidget(top_widget)
            splitter.addWidget(bottom_widget)
            splitter.setStretchFactor(0, 1)
            splitter.setStretchFactor(1, 2)
            
            # Button bar for both top and bottom sections
            button_layout = QHBoxLayout()
            
            # Command action buttons
            run_btn = QPushButton("Run Command")
            run_btn.setToolTip("Run a new command on this device")
            run_btn.clicked.connect(lambda: self._run_command_for_device(device))
            
            refresh_btn = QPushButton("Refresh")
            refresh_btn.setToolTip("Refresh the command list")
            refresh_btn.clicked.connect(lambda: self._refresh_device_commands(device_command_list))
            
            export_btn = QPushButton("Export")
            export_btn.setToolTip("Export command outputs to file")
            export_btn.clicked.connect(lambda: self._export_device_commands(device))
            
            # Control buttons for selected command
            view_btn = QPushButton("View")
            view_btn.setToolTip("View selected command in a popup window")
            view_btn.setEnabled(False)
            view_btn.clicked.connect(lambda: self._view_selected_command(device_command_list))
            
            delete_btn = QPushButton("Delete")
            delete_btn.setToolTip("Delete selected command output")
            delete_btn.setEnabled(False)
            delete_btn.clicked.connect(lambda: self._delete_selected_command(device_command_list))
            
            # Store references to enable/disable these buttons
            device_command_list.setProperty("view_button", view_btn)
            device_command_list.setProperty("delete_button", delete_btn)
            
            # Add buttons to layout
            button_layout.addWidget(run_btn)
            button_layout.addWidget(refresh_btn)
            button_layout.addWidget(export_btn)
            button_layout.addWidget(view_btn)
            button_layout.addWidget(delete_btn)
            button_layout.addStretch()
            
            # Add widgets to main layout
            self.commands_panel_widget.layout().addWidget(splitter)
            self.commands_panel_widget.layout().addLayout(button_layout)
            
            # Load commands for this device
            self._refresh_device_commands(device_command_list)
            
    def _on_command_selection_changed(self, command_list):
        """Handle command selection changed
        
        Args:
            command_list: The command list widget
        """
        selected_items = command_list.selectedItems()
        
        # Get output widgets
        output_text = command_list.property("output_text")
        output_table = command_list.property("output_table")
        output_stack = command_list.property("output_stack")
        raw_btn = command_list.property("raw_btn")
        table_btn = command_list.property("table_btn")
        
        # Enable/disable buttons based on selection
        view_btn = command_list.property("view_button")
        delete_btn = command_list.property("delete_button")
        
        if selected_items and view_btn and delete_btn:
            view_btn.setEnabled(True)
            delete_btn.setEnabled(True)
            
            # Get the selected row
            row = selected_items[0].row()
            
            # Get device, command and timestamp
            device = command_list.property("device")
            cmd_text = command_list.item(row, 0).text()
            timestamp = command_list.item(row, 1).data(Qt.UserRole)
            cmd_id = command_list.item(row, 0).data(Qt.UserRole)
            
            # Get command output
            outputs = self.get_command_outputs(device.id, cmd_id)
            
            if outputs and timestamp in outputs:
                output = outputs[timestamp]["output"]
                success = outputs[timestamp]["success"]
                
                # Update the raw text output
                if output_text:
                    output_text.setText(output)
                
                # Try to parse as table and update table view
                if output_table:
                    has_table = self._update_table_view(output, output_table)
                    
                    # Enable/disable table view based on whether we could parse a table
                    if table_btn:
                        table_btn.setEnabled(has_table)
                        
                        # If table view was selected but we can't show a table, switch to raw
                        if output_stack and output_stack.currentIndex() == 1 and not has_table:
                            output_stack.setCurrentIndex(0)
                            if raw_btn:
                                raw_btn.setChecked(True)
                            if table_btn:
                                table_btn.setChecked(False)
            else:
                # Clear outputs
                if output_text:
                    output_text.setText(f"Output not found for command: {cmd_text}")
                if output_table:
                    output_table.setRowCount(0)
                    output_table.setColumnCount(0)
                if table_btn:
                    table_btn.setEnabled(False)
        else:
            # No selection, disable buttons
            if view_btn:
                view_btn.setEnabled(False)
            if delete_btn:
                delete_btn.setEnabled(False)
                
            # Clear outputs
            if output_text:
                output_text.clear()
            if output_table:
                output_table.setRowCount(0)
                output_table.setColumnCount(0)
            if table_btn:
                table_btn.setEnabled(False)
                
    def _view_selected_command(self, command_list):
        """View the selected command in a popup window
        
        Args:
            command_list: The command list widget
        """
        selected_items = command_list.selectedItems()
        if not selected_items:
            return
            
        # Get the selected row
        row = selected_items[0].row()
        
        # Get device, command and timestamp
        device = command_list.property("device")
        cmd_text = command_list.item(row, 0).text()
        timestamp = command_list.item(row, 1).data(Qt.UserRole)
        cmd_id = command_list.item(row, 0).data(Qt.UserRole)
        
        # Get command output
        outputs = self.get_command_outputs(device.id, cmd_id)
        
        if outputs and timestamp in outputs:
            output = outputs[timestamp]["output"]
            
            # Show output in dialog
            dialog = QDialog(self.main_window)
            dialog.setWindowTitle("Command Output")
            dialog.resize(800, 600)
            
            layout = QVBoxLayout(dialog)
            
            # Add command info
            info_label = QLabel(f"Command: {cmd_text}")
            info_label.setWordWrap(True)
            layout.addWidget(info_label)
            
            # Add timestamp
            dt = datetime.datetime.fromisoformat(timestamp)
            date_str = dt.strftime("%Y-%m-%d %H:%M:%S")
            time_label = QLabel(f"Date/Time: {date_str}")
            layout.addWidget(time_label)
            
            # Add tab widget for raw/table views
            tab_widget = QTabWidget()
            
            # Raw output tab
            raw_tab = QWidget()
            raw_layout = QVBoxLayout(raw_tab)
            
            output_text = QTextEdit()
            output_text.setReadOnly(True)
            output_text.setFont(QFont("Courier New", 10))
            output_text.setText(output)
            raw_layout.addWidget(output_text)
            
            tab_widget.addTab(raw_tab, "Raw Output")
            
            # Table output tab
            table_tab = QWidget()
            table_layout = QVBoxLayout(table_tab)
            
            output_table = QTableWidget()
            output_table.setEditTriggers(QTableWidget.NoEditTriggers)
            table_layout.addWidget(output_table)
            
            # Try to parse as table
            has_table = self._update_table_view(output, output_table)
            
            # Only add table tab if we could parse it
            if has_table:
                tab_widget.addTab(table_tab, "Table View")
            
            layout.addWidget(tab_widget)
            
            # Add buttons
            buttons = QDialogButtonBox(QDialogButtonBox.Close)
            buttons.rejected.connect(dialog.reject)
            layout.addWidget(buttons)
            
            dialog.exec()
            
    def _delete_selected_command(self, command_list):
        """Delete the selected command output
        
        Args:
            command_list: The command list widget
        """
        selected_items = command_list.selectedItems()
        if not selected_items:
            return
            
        # Confirm deletion
        result = QMessageBox.question(
            self.main_window,
            "Confirm Deletion",
            "Are you sure you want to delete this command output?",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.No
        )
        
        if result != QMessageBox.Yes:
            return
            
        # Get the selected row
        row = selected_items[0].row()
        
        # Get device, command and timestamp
        device = command_list.property("device")
        cmd_id = command_list.item(row, 0).data(Qt.UserRole)
        timestamp = command_list.item(row, 1).data(Qt.UserRole)
        
        # Delete the command output
        if self.delete_command_output(device.id, cmd_id, timestamp):
            # Refresh the command list
            self._refresh_device_commands(command_list)
            
            # Show success message
            QMessageBox.information(
                self.main_window,
                "Deletion Successful",
                "Command output deleted successfully"
            )
        else:
            # Show error message
            QMessageBox.critical(
                self.main_window,
                "Deletion Failed",
                "Failed to delete command output"
            )
            
    def _run_command_for_device(self, device):
        """Run a command for the selected device
        
        Args:
            device: The device to run a command on
        """
        # Create a dialog to select and run a command
        dialog = CommandDialog(self, [device], self.main_window)
        dialog.exec()
            
    def _refresh_device_commands(self, command_list):
        """Refresh the device command list"""
        # Get output widgets
        output_text = command_list.property("output_text")
        output_table = command_list.property("output_table")
        table_btn = command_list.property("table_btn")
        
        # Clear the table
        command_list.setRowCount(0)
        
        # Clear output text if provided
        if output_text:
            output_text.clear()
            
        # Clear table view if provided
        if output_table:
            output_table.setRowCount(0)
            output_table.setColumnCount(0)
            
        # Disable table button
        if table_btn:
            table_btn.setEnabled(False)
            
        # Clear button states
        view_btn = command_list.property("view_button")
        delete_btn = command_list.property("delete_button")
        
        if view_btn:
            view_btn.setEnabled(False)
            
        if delete_btn:
            delete_btn.setEnabled(False)
        
        # Get device
        device = command_list.property("device")
        if not device:
            return
            
        # Get command outputs for device
        outputs = self.get_command_outputs(device.id)
        
        # Add rows for each command output
        for cmd_id, cmd_outputs in outputs.items():
            for timestamp, data in cmd_outputs.items():
                row = command_list.rowCount()
                command_list.insertRow(row)
                
                # Command
                cmd_text = data.get("command", cmd_id)
                cmd_item = QTableWidgetItem(cmd_text)
                cmd_item.setData(Qt.UserRole, cmd_id)
                
                # Date/time
                dt = datetime.datetime.fromisoformat(timestamp)
                dt_item = QTableWidgetItem(dt.strftime("%Y-%m-%d %H:%M:%S"))
                dt_item.setData(Qt.UserRole, timestamp)
                
                # Success
                success = "Yes" if data.get("success", True) else "No"
                success_item = QTableWidgetItem(success)
                
                # Add to row
                command_list.setItem(row, 0, cmd_item)
                command_list.setItem(row, 1, dt_item)
                command_list.setItem(row, 2, success_item)


class ReportGenerator(QDialog):
    """Dialog for generating command output reports"""
    
    def __init__(self, plugin, parent=None):
        """Initialize the dialog"""
        super().__init__(parent)
        
        self.plugin = plugin
        
        # Set dialog properties
        self.setWindowTitle("Generate Command Report")
        self.resize(700, 500)
        
        # Create UI components
        self._create_ui()
        
        # Load devices
        self._load_devices()
        
    def _create_ui(self):
        """Create the UI components"""
        # Main layout
        layout = QVBoxLayout(self)
        
        # Device selection section
        device_group = QGroupBox("Select Devices")
        device_layout = QVBoxLayout(device_group)
        
        self.device_table = QTableWidget()
        self.device_table.setColumnCount(2)
        self.device_table.setHorizontalHeaderLabels(["Device", "IP Address"])
        self.device_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.device_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeToContents)
        self.device_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.device_table.setSelectionMode(QTableWidget.MultiSelection)
        
        device_layout.addWidget(self.device_table)
        
        # Report options section
        options_group = QGroupBox("Report Options")
        options_layout = QFormLayout(options_group)
        
        # Report format
        self.report_format = QComboBox()
        self.report_format.addItems(["Text (.txt)", "HTML (.html)", "Excel (.xlsx)", "Word (.docx)"])
        options_layout.addRow("Format:", self.report_format)
        
        # Report title
        self.report_title = QLineEdit("Command Output Report")
        options_layout.addRow("Title:", self.report_title)
        
        # Include device info
        self.include_device_info = QCheckBox("Include device information")
        self.include_device_info.setChecked(True)
        options_layout.addRow("", self.include_device_info)
        
        # Include all commands
        self.include_all_commands = QCheckBox("Include all command outputs")
        self.include_all_commands.setChecked(True)
        options_layout.addRow("", self.include_all_commands)
        
        # Date range
        date_layout = QHBoxLayout()
        
        self.date_from = QLineEdit()
        self.date_from.setPlaceholderText("YYYY-MM-DD")
        
        self.date_to = QLineEdit()
        self.date_to.setPlaceholderText("YYYY-MM-DD")
        
        date_layout.addWidget(QLabel("From:"))
        date_layout.addWidget(self.date_from)
        date_layout.addWidget(QLabel("To:"))
        date_layout.addWidget(self.date_to)
        
        options_layout.addRow("Date Range:", date_layout)
        
        # Include success only
        self.success_only = QCheckBox("Include only successful commands")
        self.success_only.setChecked(False)
        options_layout.addRow("", self.success_only)
        
        # Add groups to main layout
        layout.addWidget(device_group)
        layout.addWidget(options_group)
        
        # Buttons
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        buttons.accepted.connect(self._on_generate)
        buttons.rejected.connect(self.reject)
        
        layout.addWidget(buttons)
        
    def _load_devices(self):
        """Load device list"""
        # Get all devices
        devices = self.plugin.device_manager.get_devices()
        
        # Add devices to table
        for device in devices:
            alias = device.get_property("alias", "Unnamed Device")
            ip_address = device.get_property("ip_address", "")
            
            # Skip devices without command outputs
            outputs = self.plugin.get_command_outputs(device.id)
            if not outputs:
                continue
                
            row = self.device_table.rowCount()
            self.device_table.insertRow(row)
            
            # Device info
            alias_item = QTableWidgetItem(alias)
            alias_item.setData(Qt.UserRole, device.id)
            
            ip_item = QTableWidgetItem(ip_address)
            
            # Add to table
            self.device_table.setItem(row, 0, alias_item)
            self.device_table.setItem(row, 1, ip_item)
            
    def _on_generate(self):
        """Handle generate button"""
        # Get selected devices
        selected_devices = []
        for item in self.device_table.selectedItems():
            # Make sure we only count each row once
            if item.column() == 0:
                device_id = item.data(Qt.UserRole)
                device = self.plugin.device_manager.get_device(device_id)
                if device and device not in selected_devices:
                    selected_devices.append(device)
        
        # Check if any devices are selected
        if not selected_devices:
            QMessageBox.warning(
                self,
                "No Devices Selected",
                "Please select one or more devices to include in the report."
            )
            return
            
        # Determine report format file extension
        format_index = self.report_format.currentIndex()
        if format_index == 0:
            file_ext = "txt"
            file_type = "Text Files (*.txt)"
        elif format_index == 1:
            file_ext = "html"
            file_type = "HTML Files (*.html)"
        elif format_index == 2:
            file_ext = "xlsx"
            file_type = "Excel Files (*.xlsx)"
        else:
            file_ext = "docx"
            file_type = "Word Files (*.docx)"
            
        # Get output file
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Report",
            f"{self.report_title.text().replace(' ', '_')}.{file_ext}",
            f"{file_type};;All Files (*.*)"
        )
        
        if not file_path:
            return
            
        # Validate date range if provided
        date_from = None
        date_to = None
        
        if self.date_from.text().strip():
            try:
                date_from = datetime.datetime.strptime(
                    self.date_from.text().strip(), 
                    "%Y-%m-%d"
                )
            except ValueError:
                QMessageBox.warning(
                    self,
                    "Invalid Date Format",
                    "Please enter dates in YYYY-MM-DD format."
                )
                return
                
        if self.date_to.text().strip():
            try:
                date_to = datetime.datetime.strptime(
                    self.date_to.text().strip(),
                    "%Y-%m-%d"
                )
                # Set to end of day
                date_to = date_to.replace(hour=23, minute=59, second=59)
            except ValueError:
                QMessageBox.warning(
                    self,
                    "Invalid Date Format",
                    "Please enter dates in YYYY-MM-DD format."
                )
                return
                
        # Generate the report
        try:
            if format_index == 0:
                # Text report
                self._generate_text_report(
                    file_path, 
                    selected_devices, 
                    self.report_title.text(),
                    self.include_device_info.isChecked(),
                    self.include_all_commands.isChecked(),
                    date_from,
                    date_to,
                    self.success_only.isChecked()
                )
            elif format_index == 1:
                # HTML report
                self._generate_html_report(
                    file_path, 
                    selected_devices, 
                    self.report_title.text(),
                    self.include_device_info.isChecked(),
                    self.include_all_commands.isChecked(),
                    date_from,
                    date_to,
                    self.success_only.isChecked()
                )
            elif format_index == 2:
                # Excel report
                self._generate_excel_report(
                    file_path, 
                    selected_devices, 
                    self.report_title.text(),
                    self.include_device_info.isChecked(),
                    self.include_all_commands.isChecked(),
                    date_from,
                    date_to,
                    self.success_only.isChecked()
                )
            else:
                # Word report
                self._generate_word_report(
                    file_path, 
                    selected_devices, 
                    self.report_title.text(),
                    self.include_device_info.isChecked(),
                    self.include_all_commands.isChecked(),
                    date_from,
                    date_to,
                    self.success_only.isChecked()
                )
                
            # Show success message
            QMessageBox.information(
                self,
                "Report Generated",
                f"Report saved to {file_path}"
            )
            
            # Close dialog
            self.accept()
            
        except Exception as e:
            # Show error message
            QMessageBox.critical(
                self,
                "Error Generating Report",
                f"An error occurred while generating the report: {str(e)}"
            )
            
    def _generate_text_report(self, file_path, devices, title, include_device_info, 
                             include_all, date_from, date_to, success_only):
        """Generate a text report"""
        with open(file_path, "w") as f:
            # Write report title
            f.write(f"{title}\n")
            f.write("=" * len(title) + "\n\n")
            
            # Write date
            current_date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            f.write(f"Generated: {current_date}\n\n")
            
            # Process each device
            for device in devices:
                # Write device header
                device_name = device.get_property("alias", "Unnamed Device")
                f.write(f"Device: {device_name}\n")
                f.write("-" * (len(device_name) + 8) + "\n\n")
                
                # Write device info if requested
                if include_device_info:
                    f.write("Device Information:\n")
                    for key, value in device.get_properties().items():
                        f.write(f"  {key}: {value}\n")
                    f.write("\n")
                
                # Get command outputs for this device
                outputs = self.plugin.get_command_outputs(device.id)
                
                if not outputs:
                    f.write("No command outputs available for this device.\n\n")
                    continue
                    
                # Process command outputs
                f.write("Command Outputs:\n")
                f.write("-" * 16 + "\n\n")
                
                # Track if we've written any outputs for this device
                wrote_outputs = False
                
                for cmd_id, cmd_outputs in outputs.items():
                    for timestamp, data in sorted(cmd_outputs.items()):
                        # Apply date filter if provided
                        if date_from or date_to:
                            dt = datetime.datetime.fromisoformat(timestamp)
                            
                            if date_from and dt < date_from:
                                continue
                                
                            if date_to and dt > date_to:
                                continue
                                
                        # Apply success filter if requested
                        if success_only and not data.get("success", True):
                            continue
                            
                        # Write command output
                        cmd_text = data.get("command", cmd_id)
                        output = data.get("output", "")
                        success = data.get("success", True)
                        
                        dt = datetime.datetime.fromisoformat(timestamp)
                        date_str = dt.strftime("%Y-%m-%d %H:%M:%S")
                        
                        f.write(f"Command: {cmd_text}\n")
                        f.write(f"Date/Time: {date_str}\n")
                        f.write(f"Success: {'Yes' if success else 'No'}\n")
                        f.write("Output:\n")
                        f.write("-" * 7 + "\n")
                        f.write(output + "\n\n")
                        
                        wrote_outputs = True
                        
                        # If we're not including all, just show the most recent
                        if not include_all:
                            break
                
                if not wrote_outputs:
                    f.write("No matching command outputs for this device.\n\n")
                    
                f.write("\n" + "=" * 50 + "\n\n")
                
    def _generate_html_report(self, file_path, devices, title, include_device_info, 
                             include_all, date_from, date_to, success_only):
        """Generate an HTML report"""
        with open(file_path, "w") as f:
            # Write HTML header
            f.write(f"""<!DOCTYPE html>
<html>
<head>
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        h1 {{ color: #2c3e50; }}
        h2 {{ color: #3498db; margin-top: 30px; }}
        h3 {{ color: #2980b9; }}
        pre {{ background-color: #f5f5f5; padding: 10px; border: 1px solid #ddd; overflow-x: auto; }}
        .device-info {{ background-color: #eef; padding: 10px; border: 1px solid #ddf; margin-bottom: 20px; }}
        .command {{ background-color: #efe; padding: 10px; border: 1px solid #dfd; margin-top: 20px; }}
        .command-failed {{ background-color: #fee; padding: 10px; border: 1px solid #fdd; margin-top: 20px; }}
        .timestamp {{ color: #777; font-style: italic; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p class="timestamp">Generated: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
""")
            
            # Process each device
            for device in devices:
                # Write device header
                device_name = device.get_property("alias", "Unnamed Device")
                f.write(f'    <h2>Device: {device_name}</h2>\n')
                
                # Write device info if requested
                if include_device_info:
                    f.write('    <div class="device-info">\n')
                    f.write('        <h3>Device Information</h3>\n')
                    f.write('        <table>\n')
                    
                    for key, value in device.get_properties().items():
                        f.write(f'            <tr><td><strong>{key}:</strong></td><td>{value}</td></tr>\n')
                        
                    f.write('        </table>\n')
                    f.write('    </div>\n')
                
                # Get command outputs for this device
                outputs = self.plugin.get_command_outputs(device.id)
                
                if not outputs:
                    f.write('    <p>No command outputs available for this device.</p>\n\n')
                    continue
                    
                # Process command outputs
                f.write('    <h3>Command Outputs</h3>\n')
                
                # Track if we've written any outputs for this device
                wrote_outputs = False
                
                for cmd_id, cmd_outputs in outputs.items():
                    for timestamp, data in sorted(cmd_outputs.items()):
                        # Apply date filter if provided
                        if date_from or date_to:
                            dt = datetime.datetime.fromisoformat(timestamp)
                            
                            if date_from and dt < date_from:
                                continue
                                
                            if date_to and dt > date_to:
                                continue
                                
                        # Apply success filter if requested
                        if success_only and not data.get("success", True):
                            continue
                            
                        # Write command output
                        cmd_text = data.get("command", cmd_id)
                        output = data.get("output", "")
                        success = data.get("success", True)
                        
                        dt = datetime.datetime.fromisoformat(timestamp)
                        date_str = dt.strftime("%Y-%m-%d %H:%M:%S")
                        
                        div_class = "command" if success else "command-failed"
                        
                        f.write(f'    <div class="{div_class}">\n')
                        f.write(f'        <h4>Command: {cmd_text}</h4>\n')
                        f.write(f'        <p>Date/Time: {date_str}</p>\n')
                        f.write(f'        <p>Success: {"Yes" if success else "No"}</p>\n')
                        f.write(f'        <h5>Output:</h5>\n')
                        f.write(f'        <pre>{output}</pre>\n')
                        f.write(f'    </div>\n')
                        
                        wrote_outputs = True
                        
                        # If we're not including all, just show the most recent
                        if not include_all:
                            break
                
                if not wrote_outputs:
                    f.write('    <p>No matching command outputs for this device.</p>\n\n')
                    
            # Write HTML footer
            f.write("""</body>
</html>
""")
                
    def _generate_excel_report(self, file_path, devices, title, include_device_info, 
                               include_all, date_from, date_to, success_only):
        """Generate an Excel report"""
        try:
            # Try to import openpyxl
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise RuntimeError("Excel report generation requires the openpyxl module. Please install it with 'pip install openpyxl'.")
            
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Command Report"
        
        # Set up styles
        title_font = Font(size=14, bold=True)
        header_font = Font(size=12, bold=True)
        device_font = Font(size=12, bold=True, color="0000FF")
        cmd_font = Font(size=11, bold=True)
        
        # Write title
        ws["A1"] = title
        ws["A1"].font = title_font
        ws.merge_cells("A1:G1")
        
        # Write generation timestamp
        ws["A2"] = f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        ws.merge_cells("A2:G2")
        
        row = 4
        
        # Process each device
        for device in devices:
            # Write device header
            device_name = device.get_property("alias", "Unnamed Device")
            ws[f"A{row}"] = f"Device: {device_name}"
            ws[f"A{row}"].font = device_font
            ws.merge_cells(f"A{row}:G{row}")
            row += 1
            
            # Write device info if requested
            if include_device_info:
                ws[f"A{row}"] = "Device Information:"
                ws[f"A{row}"].font = header_font
                ws.merge_cells(f"A{row}:G{row}")
                row += 1
                
                for key, value in device.get_properties().items():
                    ws[f"A{row}"] = key
                    ws[f"B{row}"] = str(value)
                    ws.merge_cells(f"B{row}:G{row}")
                    row += 1
                    
                row += 1
            
            # Get command outputs for this device
            outputs = self.plugin.get_command_outputs(device.id)
            
            if not outputs:
                ws[f"A{row}"] = "No command outputs available for this device."
                ws.merge_cells(f"A{row}:G{row}")
                row += 2
                continue
                
            # Process command outputs
            ws[f"A{row}"] = "Command Outputs:"
            ws[f"A{row}"].font = header_font
            ws.merge_cells(f"A{row}:G{row}")
            row += 1
            
            # Track if we've written any outputs for this device
            wrote_outputs = False
            
            for cmd_id, cmd_outputs in outputs.items():
                for timestamp, data in sorted(cmd_outputs.items()):
                    # Apply date filter if provided
                    if date_from or date_to:
                        dt = datetime.datetime.fromisoformat(timestamp)
                        
                        if date_from and dt < date_from:
                            continue
                            
                        if date_to and dt > date_to:
                            continue
                            
                    # Apply success filter if requested
                    if success_only and not data.get("success", True):
                        continue
                        
                    # Write command output
                    cmd_text = data.get("command", cmd_id)
                    output = data.get("output", "")
                    success = data.get("success", True)
                    
                    dt = datetime.datetime.fromisoformat(timestamp)
                    
                    # Command header
                    ws[f"A{row}"] = "Command:"
                    ws[f"B{row}"] = cmd_text
                    ws[f"B{row}"].font = cmd_font
                    ws.merge_cells(f"B{row}:G{row}")
                    row += 1
                    
                    # Date and success
                    ws[f"A{row}"] = "Date/Time:"
                    ws[f"B{row}"] = dt
                    
                    ws[f"C{row}"] = "Success:"
                    ws[f"D{row}"] = "Yes" if success else "No"
                    
                    row += 1
                    
                    # Output
                    ws[f"A{row}"] = "Output:"
                    ws.merge_cells(f"A{row}:G{row}")
                    row += 1
                    
                    # Split output into lines and write
                    output_lines = output.split("\n")
                    for line in output_lines:
                        ws[f"A{row}"] = line
                        ws.merge_cells(f"A{row}:G{row}")
                        row += 1
                        
                    row += 1
                    wrote_outputs = True
                    
                    # If we're not including all, just show the most recent
                    if not include_all:
                        break
            
            if not wrote_outputs:
                ws[f"A{row}"] = "No matching command outputs for this device."
                ws.merge_cells(f"A{row}:G{row}")
                row += 1
                
            row += 2
            
        # Adjust column widths
        for col in range(1, 8):
            ws.column_dimensions[get_column_letter(col)].width = 15
            
        # Save workbook
        wb.save(file_path)
        
    def _generate_word_report(self, file_path, devices, title, include_device_info, 
                              include_all, date_from, date_to, success_only):
        """Generate a Word report"""
        try:
            # Try to import python-docx
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError("Word report generation requires the python-docx module. Please install it with 'pip install python-docx'.")
            
        # Create document
        doc = Document()
        
        # Set title
        title_para = doc.add_heading(title, level=0)
        
        # Add generation timestamp
        timestamp_para = doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Process each device
        for device in devices:
            # Add device header
            device_name = device.get_property("alias", "Unnamed Device")
            device_para = doc.add_heading(f"Device: {device_name}", level=1)
            
            # Add device info if requested
            if include_device_info:
                doc.add_heading("Device Information", level=2)
                
                # Create a table for device properties
                property_table = doc.add_table(rows=1, cols=2)
                property_table.style = "Table Grid"
                
                # Add header row
                header_cells = property_table.rows[0].cells
                header_cells[0].text = "Property"
                header_cells[1].text = "Value"
                
                # Add properties
                for key, value in device.get_properties().items():
                    row_cells = property_table.add_row().cells
                    row_cells[0].text = key
                    row_cells[1].text = str(value)
                    
                doc.add_paragraph()
            
            # Get command outputs for this device
            outputs = self.plugin.get_command_outputs(device.id)
            
            if not outputs:
                doc.add_paragraph("No command outputs available for this device.")
                doc.add_paragraph()
                continue
                
            # Process command outputs
            doc.add_heading("Command Outputs", level=2)
            
            # Track if we've written any outputs for this device
            wrote_outputs = False
            
            for cmd_id, cmd_outputs in outputs.items():
                for timestamp, data in sorted(cmd_outputs.items()):
                    # Apply date filter if provided
                    if date_from or date_to:
                        dt = datetime.datetime.fromisoformat(timestamp)
                        
                        if date_from and dt < date_from:
                            continue
                            
                        if date_to and dt > date_to:
                            continue
                            
                    # Apply success filter if requested
                    if success_only and not data.get("success", True):
                        continue
                        
                    # Write command output
                    cmd_text = data.get("command", cmd_id)
                    output = data.get("output", "")
                    success = data.get("success", True)
                    
                    dt = datetime.datetime.fromisoformat(timestamp)
                    date_str = dt.strftime("%Y-%m-%d %H:%M:%S")
                    
                    # Command header
                    cmd_heading = doc.add_heading(level=3)
                    cmd_heading.add_run(f"Command: {cmd_text}")
                    
                    # Date and success
                    info_para = doc.add_paragraph()
                    info_para.add_run(f"Date/Time: {date_str}\n")
                    info_para.add_run(f"Success: {'Yes' if success else 'No'}")
                    
                    # Output
                    doc.add_heading("Output", level=4)
                    output_para = doc.add_paragraph()
                    output_para.style = "No Spacing"
                    output_run = output_para.add_run(output)
                    output_run.font.name = "Courier New"
                    output_run.font.size = Pt(9)
                    
                    doc.add_paragraph()
                    wrote_outputs = True
                    
                    # If we're not including all, just show the most recent
                    if not include_all:
                        break
            
            if not wrote_outputs:
                doc.add_paragraph("No matching command outputs for this device.")
                
            # Add page break after each device except the last
            if device != devices[-1]:
                doc.add_page_break()
                
        # Save document
        doc.save(file_path) 