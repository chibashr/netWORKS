#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Device importer module for NetWORKS

This module provides functionality for importing devices from various file formats 
and data sources. It is separate from the device tree to allow for better modularity.
"""

import os
import io
import csv
import json
from loguru import logger
from pathlib import Path

# Import Device class for creating device objects
from .device_manager import Device

# Try to import optional dependencies
try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False
    logger.debug("pandas not available for importing Excel files")

try:
    import xlrd
    HAS_XLRD = True
except ImportError:
    HAS_XLRD = False
    logger.debug("xlrd not available for importing legacy Excel files")

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.debug("python-docx not available for importing Word documents")

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False
    logger.debug("chardet not available for detecting file encodings")


class DeviceImporter:
    """Handles importing devices from various sources"""
    
    def __init__(self, device_manager):
        """Initialize the device importer
        
        Args:
            device_manager: The device manager instance to add imported devices to
        """
        self.device_manager = device_manager
    
    def import_from_file(self, file_path, options=None):
        """Import devices from a file
        
        Args:
            file_path: Path to the file to import
            options: Dictionary of import options:
                - delimiter: CSV delimiter character
                - has_header: Whether the file has a header row
                - encoding: File encoding (or 'auto' to detect)
                - target_group: Group to add devices to
                - field_mapping: Dictionary mapping columns to device properties
                - skip_duplicates: Whether to skip duplicate devices
                - mark_imported: Whether to tag devices as 'imported'
                
        Returns:
            tuple: (success, stats) where success is a boolean and stats is a dict with:
                - imported_count: Number of devices imported
                - skipped_count: Number of devices skipped
                - error_count: Number of devices with errors
        """
        if options is None:
            options = {}
            
        # Set default options
        options.setdefault('delimiter', ',')
        options.setdefault('has_header', True)
        options.setdefault('encoding', 'auto')
        options.setdefault('skip_duplicates', False)
        options.setdefault('mark_imported', True)
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        # Get raw data based on file type
        data, headers = self._extract_data_from_file(file_path, file_ext, options)
        
        if not data:
            logger.warning(f"No data extracted from file: {file_path}")
            return False, {"imported_count": 0, "skipped_count": 0, "error_count": 0}
            
        # Import the data using the common import function
        return self._import_data(data, headers, options)
    
    def import_from_text(self, text, options=None):
        """Import devices from pasted text
        
        Args:
            text: Text containing device data (CSV, line-by-line IP list, etc.)
            options: Dictionary of import options (see import_from_file)
                
        Returns:
            tuple: (success, stats) where success is a boolean and stats is a dict with:
                - imported_count: Number of devices imported
                - skipped_count: Number of devices skipped
                - error_count: Number of devices with errors
        """
        if options is None:
            options = {}
            
        # Set default options
        options.setdefault('delimiter', ',')
        options.setdefault('has_header', True)
        options.setdefault('skip_duplicates', False)
        options.setdefault('mark_imported', True)
        
        # Extract data from text
        data, headers = self._extract_data_from_text(text, options)
        
        if not data:
            logger.warning("No data could be extracted from text")
            return False, {"imported_count": 0, "skipped_count": 0, "error_count": 0}
        
        # Import the data using the common import function
        return self._import_data(data, headers, options)
    
    def _extract_data_from_file(self, file_path, file_ext, options):
        """Extract data from a file based on its extension
        
        Args:
            file_path: Path to the file
            file_ext: File extension (lowercase)
            options: Import options
            
        Returns:
            tuple: (data, headers) where data is a list of rows and headers is a list of column names
        """
        data = []
        headers = []
        
        try:
            # Handle different file types
            if file_ext in ['.xlsx', '.xls'] and (HAS_PANDAS or (file_ext == '.xls' and HAS_XLRD)):
                data, headers = self._extract_from_excel(file_path, file_ext, options)
            elif file_ext == '.docx' and HAS_DOCX:
                data, headers = self._extract_from_docx(file_path, options)
            else:  # CSV and text files
                data, headers = self._extract_from_csv(file_path, options)
                
            return data, headers
        except Exception as e:
            logger.error(f"Error extracting data from file: {e}", exc_info=True)
            return [], None
    
    def _extract_from_excel(self, file_path, file_ext, options):
        """Extract data from Excel files
        
        Args:
            file_path: Path to the Excel file
            file_ext: File extension (.xlsx or .xls)
            options: Import options
            
        Returns:
            tuple: (data, headers) where data is a list of rows and headers is a list of column names
        """
        has_header = options.get('has_header', True)
        
        if file_ext == '.xlsx' or (file_ext == '.xls' and HAS_PANDAS):
            # Use pandas for Excel files
            if not HAS_PANDAS:
                logger.warning("pandas is not installed, falling back to other methods")
                if file_ext == '.xls' and HAS_XLRD:
                    return self._extract_from_excel_xlrd(file_path, options)
                return [], None
                
            try:
                engine = 'xlrd' if file_ext == '.xls' else None
                df = pd.read_excel(file_path, engine=engine)
                logger.debug(f"Excel file loaded with {len(df)} rows")
                
                if has_header:
                    headers = df.columns.tolist()
                    data = df.values.tolist()
                else:
                    # Use row 0 as data, create generic headers
                    headers = [f"Column {i+1}" for i in range(len(df.columns))]
                    data = df.values.tolist()
                    
                return data, headers
            except Exception as e:
                logger.error(f"Error reading Excel file with pandas: {e}", exc_info=True)
                if file_ext == '.xls' and HAS_XLRD:
                    logger.info("Trying xlrd fallback for .xls file")
                    return self._extract_from_excel_xlrd(file_path, options)
                return [], None
        elif file_ext == '.xls' and HAS_XLRD:
            return self._extract_from_excel_xlrd(file_path, options)
            
        return [], None
    
    def _extract_from_excel_xlrd(self, file_path, options):
        """Extract data from Excel files using xlrd (for .xls files)
        
        Args:
            file_path: Path to the Excel file
            options: Import options
            
        Returns:
            tuple: (data, headers) where data is a list of rows and headers is a list of column names
        """
        has_header = options.get('has_header', True)
        
        try:
            # Use xlrd directly
            workbook = xlrd.open_workbook(file_path)
            sheet = workbook.sheet_by_index(0)
            logger.debug(f"XLS file loaded with {sheet.nrows} rows")
            
            # Get all rows
            all_rows = [sheet.row_values(i) for i in range(sheet.nrows)]
            
            if has_header and len(all_rows) > 0:
                headers = all_rows[0]
                data = all_rows[1:]
            else:
                headers = [f"Column {i+1}" for i in range(sheet.ncols)]
                data = all_rows
                
            return data, headers
        except Exception as e:
            logger.error(f"Error reading Excel file with xlrd: {e}", exc_info=True)
            return [], None
    
    def _extract_from_docx(self, file_path, options):
        """Extract data from Word documents
        
        Args:
            file_path: Path to the Word document
            options: Import options
            
        Returns:
            tuple: (data, headers) where data is a list of rows and headers is a list of column names
        """
        try:
            doc = Document(file_path)
            
            # Try to find tables
            if doc.tables:
                # Get first table
                table = doc.tables[0]
                
                # Extract rows from table
                has_header = options.get('has_header', True)
                all_rows = []
                
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    all_rows.append(row_data)
                
                if not all_rows:
                    logger.warning("No data found in Word table")
                    return [], None
                    
                if has_header and len(all_rows) > 0:
                    headers = all_rows[0]
                    data = all_rows[1:]
                else:
                    headers = [f"Column {i+1}" for i in range(len(all_rows[0]))]
                    data = all_rows
                    
                return data, headers
            else:
                # No tables, try to extract from paragraphs
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                return self._extract_data_from_text(text, options)
                
        except Exception as e:
            logger.error(f"Error extracting data from Word document: {e}", exc_info=True)
            return [], None
            
    def _extract_from_csv(self, file_path, options):
        """Extract data from CSV files
        
        Args:
            file_path: Path to the CSV file
            options: Import options
            
        Returns:
            tuple: (data, headers) where data is a list of rows and headers is a list of column names
        """
        # Get selected encoding or detect
        encoding_option = options.get('encoding', 'auto')
        encoding = None
        
        if encoding_option == 'auto':
            # Detect encoding if chardet is available
            if HAS_CHARDET:
                try:
                    with open(file_path, 'rb') as f:
                        raw_data = f.read(10000)  # Read first 10000 bytes
                        result = chardet.detect(raw_data)
                        encoding = result['encoding']
                        logger.debug(f"Detected encoding: {encoding} (confidence: {result.get('confidence', 0):.2f})")
                        if not encoding:
                            encoding = 'utf-8'  # Fallback to UTF-8
                except Exception as e:
                    logger.error(f"Error detecting file encoding: {e}")
                    encoding = 'utf-8'  # Fallback to UTF-8
            else:
                encoding = 'utf-8'  # Fallback to UTF-8 if chardet not available
        else:
            encoding = encoding_option
            
        delimiter = options.get('delimiter', ',')
        has_header = options.get('has_header', True)
        
        # Handle different delimiter options
        delimiter_map = {
            "Comma (,)": ',',
            "Tab": '\t',
            "Semicolon (;)": ';',
            "Pipe (|)": '|',
            "Space": ' '
        }
        if isinstance(delimiter, str) and delimiter in delimiter_map:
            delimiter = delimiter_map[delimiter]
        
        # Read the file with detected encoding
        try:
            with open(file_path, 'r', newline='', encoding=encoding, errors='replace') as f:
                reader = csv.reader(f, delimiter=delimiter)
                rows = list(reader)
                
                if not rows:
                    logger.warning(f"CSV file is empty: {file_path}")
                    return [], None
                
                logger.debug(f"CSV file loaded with {len(rows)} rows")
                
                if has_header:
                    headers = rows[0]
                    data = rows[1:]
                else:
                    if not rows[0]:
                        logger.warning(f"First row is empty in CSV file: {file_path}")
                        return [], None
                        
                    headers = [f"Column {i+1}" for i in range(len(rows[0]))]
                    data = rows
                    
                return data, headers
        except csv.Error as e:
            logger.error(f"CSV parsing error: {e}")
            return [], None
        except Exception as e:
            logger.error(f"File reading error: {e}")
            return [], None
    
    def _extract_data_from_text(self, text, options):
        """Extract data from pasted text
        
        Args:
            text: Text containing device data
            options: Import options
            
        Returns:
            tuple: (data, headers) where data is a list of rows and headers is a list of column names
        """
        if not text or not text.strip():
            logger.warning("Empty text provided for import")
            return [], None
            
        # Get options
        delimiter = options.get('delimiter', ',')
        has_header = options.get('has_header', True)
        
        # Handle different delimiter options
        delimiter_map = {
            "Comma (,)": ',',
            "Tab": '\t',
            "Semicolon (;)": ';',
            "Pipe (|)": '|',
            "Space": ' '
        }
        if isinstance(delimiter, str) and delimiter in delimiter_map:
            delimiter = delimiter_map[delimiter]
            
        # Check if this is a simple list (one entry per line)
        lines = text.splitlines()
        
        # Detect if it's a simple list of IPs/hostnames or CSV data
        is_simple_list = True
        for line in lines[:10]:  # Check first 10 lines
            if line.strip() and delimiter in line:
                is_simple_list = False
                break
        
        # Handle simple list of IPs or hostnames (one per line)
        if is_simple_list and len(lines) > 0:
            logger.debug("Detected simple list - one entry per line")
            valid_lines = [line.strip() for line in lines if line.strip()]
            if valid_lines:
                # Create a simple csv structure with one column
                data = [[line] for line in valid_lines]
                headers = ["ip_address"]
                logger.debug(f"Created {len(data)} rows with 1 column from line-by-line input")
                return data, headers
        
        # Otherwise process as normal CSV
        try:
            logger.debug(f"Processing as CSV with delimiter: '{delimiter}'")
            f = io.StringIO(text)
            reader = csv.reader(f, delimiter=delimiter)
            rows = list(reader)
            
            if not rows:
                logger.warning("No rows found in pasted text")
                return [], None
            
            # Check if we have any non-empty rows
            valid_rows = []
            for row in rows:
                if row and any(cell.strip() for cell in row if cell):
                    valid_rows.append(row)
            
            if not valid_rows:
                logger.warning("No valid data rows found in pasted text")
                return [], None
            
            # Use valid rows for further processing
            rows = valid_rows
            
            if has_header and len(rows) > 1:
                headers = rows[0]
                data = rows[1:]
            else:
                headers = [f"Column {i+1}" for i in range(len(rows[0]))]
                data = rows
                
            return data, headers
        except Exception as e:
            logger.error(f"Error parsing CSV text: {e}", exc_info=True)
            return [], None
    
    def _import_data(self, data, headers, options):
        """Import device data from extracted rows
        
        Args:
            data: List of data rows
            headers: List of column headers
            options: Import options
            
        Returns:
            tuple: (success, stats) where success is a boolean and stats is a dict with:
                - imported_count: Number of devices imported
                - skipped_count: Number of devices skipped
                - error_count: Number of devices with errors
        """
        if not data or not headers:
            return False, {"imported_count": 0, "skipped_count": 0, "error_count": 0}
            
        # Initialize stats
        stats = {
            "imported_count": 0,
            "skipped_count": 0,
            "error_count": 0
        }
        
        # Get options
        field_mapping = options.get('field_mapping', {})
        skip_duplicates = options.get('skip_duplicates', False)
        mark_imported = options.get('mark_imported', True)
        target_group = options.get('target_group', None)
        
        # If no field mapping provided, try to auto-detect
        if not field_mapping:
            field_mapping = self._auto_detect_field_mapping(headers)
            
        # Check for duplicates if needed
        existing_ips = {}
        existing_hostnames = {}
        
        if skip_duplicates:
            for device in self.device_manager.get_devices():
                ip = device.get_property("ip_address")
                hostname = device.get_property("hostname")
                
                if ip and ip.strip():
                    existing_ips[ip.strip()] = device
                if hostname and hostname.strip():
                    existing_hostnames[hostname.strip()] = device
        
        # Process each row and create devices
        for row_data in data:
            # Skip empty rows or rows with only empty strings
            if not row_data:
                continue
                
            # Check if all cells are empty strings (if they're strings)
            if all((isinstance(cell, str) and cell.strip() == "") for cell in row_data if cell is not None):
                continue
                
            # Create base device properties
            device_props = {}
            
            # Map fields based on the field mapping
            for i, header in enumerate(headers):
                if i < len(row_data):
                    value = row_data[i]
                    
                    # Find the device property for this header
                    prop_name = None
                    
                    # Check if the header is in field_mapping
                    for field_name, mapped_headers in field_mapping.items():
                        if header in mapped_headers:
                            prop_name = field_name
                            break
                    
                    # If not found in mapping, check for exact matches with common property names
                    if prop_name is None:
                        header_lower = header.lower()
                        common_mappings = {
                            "alias": ["name", "device name", "alias", "hostname", "host"],
                            "hostname": ["hostname", "host", "host name", "device name"],
                            "ip_address": ["ip", "ip address", "ipaddress", "address"],
                            "mac_address": ["mac", "mac address", "macaddress", "physical", "physical address"],
                            "notes": ["notes", "description", "comments"],
                            "tags": ["tags", "labels", "categories"]
                        }
                        
                        for prop, aliases in common_mappings.items():
                            if header_lower in aliases:
                                prop_name = prop
                                break
                    
                    # If still not found, use the header as the property name
                    if prop_name is None:
                        prop_name = header
                    
                    # Convert value to appropriate type if needed
                    if value not in (None, ""):
                        # Handle tags as a list
                        if prop_name == "tags" and isinstance(value, str):
                            tags = [tag.strip() for tag in value.split(",") if tag.strip()]
                            device_props[prop_name] = tags
                        else:
                            device_props[prop_name] = value
            
            # Skip if we don't have either IP address or hostname
            if not device_props.get("ip_address") and not device_props.get("hostname"):
                logger.debug(f"Skipping row with no IP or hostname: {row_data}")
                stats["skipped_count"] += 1
                continue
                
            # Check for duplicates
            if skip_duplicates:
                ip = device_props.get("ip_address", "").strip()
                hostname = device_props.get("hostname", "").strip()
                
                if ip and ip in existing_ips:
                    logger.debug(f"Skipping duplicate IP: {ip}")
                    stats["skipped_count"] += 1
                    continue
                    
                if hostname and hostname in existing_hostnames:
                    logger.debug(f"Skipping duplicate hostname: {hostname}")
                    stats["skipped_count"] += 1
                    continue
            
            # Add imported tag if option is selected
            if mark_imported:
                tags = device_props.get("tags", [])
                if isinstance(tags, list):
                    if "imported" not in tags:
                        tags.append("imported")
                else:
                    tags = [tags, "imported"] if tags else ["imported"]
                device_props["tags"] = tags
                
            try:
                # Ensure all device properties have valid types
                # Extract alias first with a default value
                alias = device_props.pop("alias", "Imported Device")
                
                # Ensure string properties are strings
                for str_prop in ["hostname", "ip_address", "mac_address", "notes", "status"]:
                    if str_prop in device_props:
                        # Convert to string if not None
                        if device_props[str_prop] is not None:
                            device_props[str_prop] = str(device_props[str_prop])
                        else:
                            device_props[str_prop] = ""
                
                # Ensure tags is a list
                if "tags" in device_props and not isinstance(device_props["tags"], list):
                    if device_props["tags"] is not None:
                        device_props["tags"] = [str(device_props["tags"])]
                    else:
                        device_props["tags"] = []
                
                # Create a device with minimal valid properties
                device = Device(alias=alias, **device_props)
                
                # Add to device manager
                self.device_manager.add_device(device)
                
                # Add to group if specified
                if target_group and target_group != self.device_manager.root_group:
                    self.device_manager.add_device_to_group(device, target_group)
                
                # Update tracking for duplicates
                if skip_duplicates:
                    ip = device_props.get("ip_address", "").strip()
                    hostname = device_props.get("hostname", "").strip()
                    if ip:
                        existing_ips[ip] = device
                    if hostname:
                        existing_hostnames[hostname] = device
                
                stats["imported_count"] += 1
                
            except Exception as e:
                logger.error(f"Error creating device: {e}", exc_info=True)
                stats["error_count"] += 1
        
        return stats["imported_count"] > 0, stats
    
    def _auto_detect_field_mapping(self, headers):
        """Auto-detect field mappings based on headers"""
        mappings = {}
        
        for header in headers:
            header_lower = header.lower()
            
            # Default to custom
            mapping = "custom"
            
            # Check common mappings
            if header_lower in ["name", "device name", "alias", "device"]:
                mapping = "alias"
            elif header_lower in ["hostname", "host", "host name"]:
                mapping = "hostname"
            elif header_lower in ["ip", "ip address", "ipaddress", "address", "ipv4", "ipv4 address"]:
                mapping = "ip_address"
            elif header_lower in ["mac", "mac address", "macaddress", "physical address", "physical"]:
                mapping = "mac_address"
            elif header_lower in ["status", "state"]:
                mapping = "status"
            elif header_lower in ["notes", "description", "comments", "comment"]:
                mapping = "notes"
            elif header_lower in ["tags", "labels", "categories", "category"]:
                mapping = "tags"
            elif header_lower in ["vendor", "manufacturer"]:
                mapping = "vendor"
            elif header_lower in ["model", "device model"]:
                mapping = "model"
            elif header_lower in ["serial", "serial number", "serialnumber"]:
                mapping = "serial_number"
            elif header_lower in ["location", "site", "building", "room"]:
                mapping = "location"
            
            # Add to mappings
            if mapping not in mappings:
                mappings[mapping] = []
            mappings[mapping].append(header)
        
        return mappings

    def run_import_wizard(self, parent=None):
        """Run the import wizard UI
        
        Args:
            parent: Parent widget
            
        Returns:
            bool: True if import was successful, False otherwise
        """
        # Import Qt modules here to avoid circular imports
        from PySide6.QtWidgets import (
            QWizard, QWizardPage, QVBoxLayout, QHBoxLayout, QFormLayout,
            QGroupBox, QLabel, QLineEdit, QPushButton, QRadioButton, 
            QButtonGroup, QCheckBox, QComboBox, QTableWidget, 
            QTableWidgetItem, QAbstractItemView, QPlainTextEdit,
            QFileDialog, QMessageBox, QDialog
        )
        from PySide6.QtCore import Qt, Signal, QModelIndex
        import os
        
        # Create the wizard dialog
        wizard = QWizard(parent)
        wizard.setWindowTitle("Import Devices")
        wizard.setMinimumSize(800, 600)
        
        # First page - choose file or paste text
        intro_page = QWizardPage()
        intro_page.setTitle("Import Devices")
        intro_page.setSubTitle("Choose import method")
        
        intro_layout = QVBoxLayout(intro_page)
        
        method_group = QButtonGroup(intro_page)
        file_radio = QRadioButton("Import from file")
        text_radio = QRadioButton("Paste text data")
        
        file_radio.setChecked(True)
        method_group.addButton(file_radio)
        method_group.addButton(text_radio)
        
        intro_layout.addWidget(file_radio)
        intro_layout.addWidget(text_radio)
        
        # Add page and store page ID
        intro_page_id = wizard.addPage(intro_page)
        
        # Second page - file selection (only shown if file radio is selected)
        file_page = QWizardPage()
        file_page.setTitle("File Import")
        file_page.setSubTitle("Select a file to import devices from")
        
        file_layout = QVBoxLayout(file_page)
        
        file_select_layout = QHBoxLayout()
        file_path_edit = QLineEdit()
        file_path_edit.setPlaceholderText("Select a file...")
        file_select_layout.addWidget(file_path_edit)
        
        def select_file():
            filter_str = "CSV Files (*.csv);;Text Files (*.txt)"
            if HAS_PANDAS:
                filter_str += ";;Excel Files (*.xlsx *.xls)"
            if HAS_XLRD:
                filter_str += ";;Excel Files (*.xls)"
            
            filter_str = "All Supported Files (" + \
                         "*.csv *.txt" + \
                         (" *.xlsx *.xls" if HAS_PANDAS else "") + \
                         (" *.xls" if HAS_XLRD else "") + \
                         ");;" + filter_str
            
            file_path, _ = QFileDialog.getOpenFileName(
                wizard, "Select Import File", "", filter_str
            )
            if file_path:
                file_path_edit.setText(file_path)
                logger.debug(f"Import file selected: {file_path}")
                # Make sure the Next button is enabled
                wizard.button(QWizard.NextButton).setEnabled(True)
        
        browse_button = QPushButton("Browse...")
        browse_button.clicked.connect(select_file)
        file_select_layout.addWidget(browse_button)
        file_layout.addLayout(file_select_layout)
        
        file_options_group = QGroupBox("File Options")
        file_options_layout = QFormLayout(file_options_group)
        
        delimiter_combo = QComboBox()
        delimiter_combo.addItems(["Comma (,)", "Tab", "Semicolon (;)", "Pipe (|)", "Space"])
        file_options_layout.addRow("Delimiter:", delimiter_combo)
        
        has_header_check = QCheckBox("First row contains headers")
        has_header_check.setChecked(True)
        file_options_layout.addRow("", has_header_check)
        
        encoding_combo = QComboBox()
        encoding_combo.addItems(["Auto-detect", "UTF-8", "ASCII", "Latin-1 (ISO-8859-1)", "Windows-1252"])
        file_options_layout.addRow("Text Encoding:", encoding_combo)
        
        file_layout.addWidget(file_options_group)
        
        # Add page validation
        def file_page_isComplete():
            # Check if a file is selected
            return bool(file_path_edit.text())
        
        file_page.isComplete = file_page_isComplete
        file_page_id = wizard.addPage(file_page)
        
        # Connect file path changes to update the Next button state
        def on_file_path_changed():
            # Enable Next button when a file path is entered
            has_file = bool(file_path_edit.text())
            wizard.button(QWizard.NextButton).setEnabled(has_file)
            
        file_path_edit.textChanged.connect(on_file_path_changed)
        
        # Initial state - disable Next button if no file selected
        wizard.button(QWizard.NextButton).setEnabled(False)
        
        # Text page (only shown if text radio is selected)
        text_page = QWizardPage()
        text_page.setTitle("Text Import")
        text_page.setSubTitle("Paste data below")
        
        text_layout = QVBoxLayout(text_page)
        
        text_options_group = QGroupBox("Text Options")
        text_options_layout = QFormLayout(text_options_group)
        
        text_delimiter_combo = QComboBox()
        text_delimiter_combo.addItems(["Comma (,)", "Tab", "Semicolon (;)", "Pipe (|)", "Space", "Auto-detect"])
        text_options_layout.addRow("Delimiter:", text_delimiter_combo)
        
        text_has_header_check = QCheckBox("First row contains headers")
        text_has_header_check.setChecked(True)
        text_options_layout.addRow("", text_has_header_check)
        
        text_layout.addWidget(text_options_group)
        
        text_edit = QPlainTextEdit()
        text_edit.setPlaceholderText("Paste data here...")
        text_layout.addWidget(text_edit)
        
        # Help text
        help_label = QLabel("You can paste CSV data, or a simple list of IP addresses or hostnames (one per line)")
        help_label.setWordWrap(True)
        text_layout.addWidget(help_label)
        
        # Add page validation
        def text_page_isComplete():
            # Check if text is entered
            return bool(text_edit.toPlainText().strip())
        
        text_page.isComplete = text_page_isComplete
        text_page_id = wizard.addPage(text_page)
        
        # Connect text edit changes to completeChanged signal
        def on_text_changed():
            # Enable the Next button when there's text
            has_text = bool(text_edit.toPlainText().strip())
            wizard.button(QWizard.NextButton).setEnabled(has_text)
        
        text_edit.textChanged.connect(on_text_changed)
        
        # Initial state - disable Next button if no text
        wizard.button(QWizard.NextButton).setEnabled(False)
        
        # Field mapping page
        mapping_page = QWizardPage()
        mapping_page.setTitle("Field Mapping")
        mapping_page.setSubTitle("Map columns to device properties")
        
        mapping_layout = QVBoxLayout(mapping_page)
        
        # Mapping table
        mapping_table = QTableWidget()
        mapping_table.setColumnCount(2)
        mapping_table.setHorizontalHeaderLabels(["Field", "Device Property"])
        mapping_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        mapping_table.horizontalHeader().setStretchLastSection(True)
        
        mapping_layout.addWidget(mapping_table)
        
        # Preview table
        preview_group = QGroupBox("Data Preview")
        preview_layout = QVBoxLayout(preview_group)
        
        preview_table = QTableWidget()
        preview_layout.addWidget(preview_table)
        
        mapping_layout.addWidget(preview_group)
        
        mapping_page_id = wizard.addPage(mapping_page)
        
        # Options page
        options_page = QWizardPage()
        options_page.setTitle("Import Options")
        options_page.setSubTitle("Set additional import options")
        
        options_layout = QVBoxLayout(options_page)
        
        target_group_label = QLabel("Target Group:")
        options_layout.addWidget(target_group_label)
        
        target_group_combo = QComboBox()
        target_group_combo.addItem("All Devices")
        
        # Add all groups
        for group in self.device_manager.get_groups():
            if group != self.device_manager.root_group:
                target_group_combo.addItem(group.name)
        
        options_layout.addWidget(target_group_combo)
        
        # New group option
        new_group_layout = QHBoxLayout()
        new_group_check = QCheckBox("Create new group:")
        new_group_edit = QLineEdit()
        new_group_edit.setEnabled(False)
        
        def toggle_new_group():
            new_group_edit.setEnabled(new_group_check.isChecked())
            if new_group_check.isChecked():
                target_group_combo.setEnabled(False)
            else:
                target_group_combo.setEnabled(True)
        
        new_group_check.toggled.connect(toggle_new_group)
        
        new_group_layout.addWidget(new_group_check)
        new_group_layout.addWidget(new_group_edit)
        options_layout.addLayout(new_group_layout)
        
        # Options
        skip_duplicates_check = QCheckBox("Skip duplicates (based on IP address and hostname)")
        skip_duplicates_check.setChecked(True)
        options_layout.addWidget(skip_duplicates_check)
        
        mark_imported_check = QCheckBox("Add 'imported' tag to devices")
        mark_imported_check.setChecked(True)
        options_layout.addWidget(mark_imported_check)
        
        options_page_id = wizard.addPage(options_page)
        
        # Confirmation page
        confirm_page = QWizardPage()
        confirm_page.setTitle("Confirm Import")
        confirm_page.setSubTitle("Ready to import devices")
        
        confirm_layout = QVBoxLayout(confirm_page)
        
        summary_label = QLabel()
        summary_label.setWordWrap(True)
        confirm_layout.addWidget(summary_label)
        
        confirm_page_id = wizard.addPage(confirm_page)
        
        # Prepare data function
        def prepare_data():
            data = []
            headers = None
            
            try:
                if file_radio.isChecked():
                    # File import
                    file_path = file_path_edit.text()
                    if not file_path or not os.path.exists(file_path):
                        logger.error(f"File not found: {file_path}")
                        return [], None
                    
                    # Collect options
                    options = {
                        'delimiter': delimiter_combo.currentText(),
                        'has_header': has_header_check.isChecked(),
                        'encoding': encoding_combo.currentText(),
                    }
                    
                    # Extract data from the file
                    data, headers = self._extract_data_from_file(
                        file_path, 
                        os.path.splitext(file_path)[1].lower(), 
                        options
                    )
                else:
                    # Text import
                    text = text_edit.toPlainText()
                    if not text.strip():
                        logger.error("No text provided")
                        return [], None
                    
                    # Collect options
                    options = {
                        'delimiter': text_delimiter_combo.currentText(),
                        'has_header': text_has_header_check.isChecked(),
                    }
                    
                    # Extract data from text
                    data, headers = self._extract_data_from_text(text, options)
                
                return data, headers
            except Exception as e:
                logger.error(f"Error preparing data: {e}", exc_info=True)
                QMessageBox.critical(wizard, "Import Error", f"Error preparing data: {str(e)}")
                return [], None
        
        # Wizard page navigation
        def on_current_id_changed(page_id):
            if page_id == file_page_id:
                # Only show file page if file import is selected
                if not file_radio.isChecked():
                    wizard.next()
            elif page_id == text_page_id:
                # Only show text page if text import is selected
                if not text_radio.isChecked():
                    wizard.next()
            elif page_id == mapping_page_id:
                # Populate mapping page
                data, headers = prepare_data()
                
                if not data or not headers:
                    QMessageBox.critical(wizard, "Import Error", "No data could be extracted")
                    wizard.back()
                    return
                
                # Update mapping table
                mapping_table.setRowCount(len(headers))
                
                # Available device properties
                properties = [
                    "alias", "hostname", "ip_address", "mac_address", 
                    "status", "notes", "tags", "vendor", "model", 
                    "serial_number", "location", "custom"
                ]
                
                # Auto-detect mappings based on header names
                auto_mappings = self._auto_detect_field_mapping(headers)
                
                # Add headers to the mapping table
                for i, header in enumerate(headers):
                    # Add the header
                    header_item = QTableWidgetItem(header)
                    header_item.setFlags(header_item.flags() & ~Qt.ItemIsEditable)
                    mapping_table.setItem(i, 0, header_item)
                    
                    # Add combo box for mapping
                    combo = QComboBox()
                    combo.addItems(properties)
                    
                    # Find the mapping for this header
                    for field_name, mapped_headers in auto_mappings.items():
                        if header in mapped_headers:
                            combo.setCurrentText(field_name)
                            break
                    
                    mapping_table.setCellWidget(i, 1, combo)
                
                # Update preview table
                max_rows = min(5, len(data))
                preview_table.setRowCount(max_rows)
                preview_table.setColumnCount(len(headers))
                preview_table.setHorizontalHeaderLabels(headers)
                
                for row in range(max_rows):
                    for col in range(len(headers)):
                        if col < len(data[row]):
                            item = QTableWidgetItem(str(data[row][col]))
                            preview_table.setItem(row, col, item)
            
            elif page_id == confirm_page_id:
                # Create summary text for confirmation page
                data, headers = prepare_data()
                row_count = len(data) if data else 0
                
                if new_group_check.isChecked():
                    target = f"Create new group: {new_group_edit.text()}"
                else:
                    target = f"Add to group: {target_group_combo.currentText()}"
                    
                summary = f"""
                Ready to import {row_count} device(s).
                
                {target}
                
                Skip duplicates: {"Yes" if skip_duplicates_check.isChecked() else "No"}
                Add 'imported' tag: {"Yes" if mark_imported_check.isChecked() else "No"}
                """
                
                summary_label.setText(summary)
        
        wizard.currentIdChanged.connect(on_current_id_changed)
        
        # Import function
        def import_devices():
            try:
                # Get the data
                data, headers = prepare_data()
                
                if not data:
                    logger.error("No data to import")
                    QMessageBox.critical(
                        wizard,
                        "Import Error",
                        "No data to import"
                    )
                    return False
                
                # Get field mapping
                field_mapping = {}
                for i in range(mapping_table.rowCount()):
                    header = mapping_table.item(i, 0).text()
                    mapping = mapping_table.cellWidget(i, 1).currentText()
                    
                    # Add to mapping
                    if mapping not in field_mapping:
                        field_mapping[mapping] = []
                    field_mapping[mapping].append(header)
                
                # Get or create target group
                target_group = None
                if new_group_check.isChecked():
                    group_name = new_group_edit.text()
                    if group_name:
                        target_group = self.device_manager.create_group(group_name)
                else:
                    group_name = target_group_combo.currentText()
                    target_group = self.device_manager.get_group(group_name)
                
                # Prepare import options
                import_options = {
                    'field_mapping': field_mapping,
                    'skip_duplicates': skip_duplicates_check.isChecked(),
                    'mark_imported': mark_imported_check.isChecked(),
                    'target_group': target_group
                }
                
                # Perform the import
                if file_radio.isChecked():
                    # File import
                    file_path = file_path_edit.text()
                    import_options['delimiter'] = delimiter_combo.currentText()
                    import_options['has_header'] = has_header_check.isChecked()
                    import_options['encoding'] = encoding_combo.currentText()
                    
                    success, stats = self.import_from_file(file_path, import_options)
                else:
                    # Text import
                    text = text_edit.toPlainText()
                    import_options['delimiter'] = text_delimiter_combo.currentText()
                    import_options['has_header'] = text_has_header_check.isChecked()
                    
                    success, stats = self.import_from_text(text, import_options)
                
                # Show results
                if success:
                    QMessageBox.information(
                        wizard,
                        "Import Successful",
                        f"Successfully imported {stats['imported_count']} device(s).\n"
                        f"Skipped: {stats['skipped_count']}\n"
                        f"Errors: {stats['error_count']}"
                    )
                    return True
                else:
                    QMessageBox.warning(
                        wizard,
                        "Import Warning",
                        f"Import completed with warnings.\n"
                        f"Imported: {stats['imported_count']}\n"
                        f"Skipped: {stats['skipped_count']}\n"
                        f"Errors: {stats['error_count']}"
                    )
                    return stats['imported_count'] > 0
            except Exception as e:
                logger.error(f"Error during import: {e}", exc_info=True)
                QMessageBox.critical(
                    wizard,
                    "Import Error",
                    f"An error occurred during import: {str(e)}"
                )
                return False
        
        # Run the wizard
        if wizard.exec_() == QDialog.Accepted:
            return import_devices()
        
        return False 